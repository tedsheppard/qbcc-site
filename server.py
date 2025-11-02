import os, re, shutil, sqlite3, requests, unicodedata, pandas as pd, io, json, sys
from urllib.parse import unquote_plus
from fastapi import FastAPI, Query, Form, Path, HTTPException, UploadFile, File, Body
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from email.message import EmailMessage
import aiosmtplib
from openai import OpenAI
from google.cloud import storage
from apscheduler.schedulers.background import BackgroundScheduler
from datetime import datetime, timedelta, date
from jose import JWTError, jwt
import PyPDF2
import docx
import extract_msg
import time
import random
import string
import io
import requests
import time
import os
import bcrypt
import pypandoc
from typing import List, Optional
import zipstream
import pytesseract
from PIL import Image
import stripe
import secrets
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi import Depends  # Add this if not already imported
from io import BytesIO
from datetime import datetime
from fastapi import Response
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from fastapi import Header, Depends

def get_gcs_client():
    """
    Creates a GCS client by reading JSON credentials from an environment variable,
    writing them to a temporary file, and authenticating with that file.
    """
    # Step 1: Read the JSON content from our new environment variable.
    gcs_credentials_json = os.getenv("GCS_CREDENTIALS_JSON")
    
    if not gcs_credentials_json:
        print("FATAL: GCS_CREDENTIALS_JSON environment variable not found.", file=sys.stderr)
        return None

    # Step 2: Define a path for our temporary credentials file inside the Render container.
    temp_credentials_path = "/tmp/gcs_credentials.json"

    try:
        # Step 3: Write the JSON content to the temporary file.
        with open(temp_credentials_path, "w") as f:
            f.write(gcs_credentials_json)
        
        print(f"Successfully wrote GCS credentials to {temp_credentials_path}")

        # Step 4: Authenticate using the temporary file we just created.
        client = storage.Client.from_service_account_json(temp_credentials_path)
        print("GCS client created successfully from temporary file.")
        return client
        
    except Exception as e:
        print(f"FATAL: Failed to create GCS client from credentials file. Error: {e}", file=sys.stderr)
        return None
        
# ---------------- setup ----------------
ROOT = os.path.dirname(os.path.abspath(__file__))
SITE_DIR = os.path.join(ROOT, "site")
DB_PATH = "/tmp/qbcc.db"
LEXIFILE_STORAGE = "/tmp/lexifile_storage"
os.makedirs(LEXIFILE_STORAGE, exist_ok=True)


# Download DB from GCS on startup if it doesn't exist in /tmp
if not os.path.exists(DB_PATH):
    try:
        gcs_bucket_name = os.getenv("GCS_BUCKET_NAME")
        gcs_db_object_name = os.getenv("GCS_DB_OBJECT_NAME", "qbcc.db")
        
        if gcs_bucket_name:
            print(f"Database not found at {DB_PATH}. Downloading from GCS bucket '{gcs_bucket_name}'...")
            storage_client = get_gcs_client()
            bucket = storage_client.bucket(gcs_bucket_name)
            blob = bucket.blob(gcs_db_object_name)
            blob.download_to_filename(DB_PATH)
            print("Database downloaded successfully.")
        else:
            # Fallback for local development if GCS env var isn't set
            print("GCS_BUCKET_NAME not set. Trying to copy local 'qbcc.db'.")
            if os.path.exists("qbcc.db"):
                shutil.copy("qbcc.db", DB_PATH)
            else:
                print("FATAL: No local 'qbcc.db' found and GCS bucket not configured.")
    except Exception as e:
        print(f"FATAL: Failed to download database from GCS. Error: {e}")

# sqlite connection
con = sqlite3.connect(DB_PATH, check_same_thread=False)
con.row_factory = sqlite3.Row
con.execute("PRAGMA cache_size = -20000")
con.execute("PRAGMA temp_store = MEMORY")
con.execute("PRAGMA mmap_size = 30000000000")
con.execute("PRAGMA journal_mode=WAL;")


app = FastAPI()

# --- UNIFIED USERS DATABASE CONNECTION ---
PURCHASES_DB_PATH = "/var/data/adjudicator_purchases.db"
purchases_con = sqlite3.connect(PURCHASES_DB_PATH, check_same_thread=False)
purchases_con.row_factory = sqlite3.Row
purchases_cur = purchases_con.cursor()

# In server.py, near the top where you define purchases_con

purchases_cur = purchases_con.cursor()

# --- Core user accounts ---
purchases_cur.execute("""
CREATE TABLE IF NOT EXISTS purchase_users (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    email TEXT UNIQUE NOT NULL,
    hashed_password TEXT NOT NULL,
    first_name TEXT,
    last_name TEXT,
    firm_name TEXT,
    abn TEXT,
    billing_address TEXT,
    billing_city TEXT,
    billing_state TEXT,
    billing_postcode TEXT,
    employee_size TEXT,
    phone TEXT,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    last_login TIMESTAMP,
    email_verified BOOLEAN DEFAULT 0
)
""")

# In server.py, add this with your other table creation statements

# --- Full Access Users ---
purchases_cur.execute("""
CREATE TABLE IF NOT EXISTS full_access_users (
    email TEXT PRIMARY KEY NOT NULL,
    added_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
)
""")
purchases_con.commit()

# --- Purchase history ---
purchases_cur.execute("""
CREATE TABLE IF NOT EXISTS adjudicator_purchases (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    user_email TEXT NOT NULL,
    adjudicator_name TEXT NOT NULL,
    purchase_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    stripe_payment_intent_id TEXT,
    amount_paid REAL DEFAULT 69.95,
    stripe_invoice_id TEXT,
    UNIQUE(user_email, adjudicator_name)
)
""")

# --- Password reset tokens ---
purchases_cur.execute("""
CREATE TABLE IF NOT EXISTS password_reset_tokens (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    email TEXT NOT NULL,
    token TEXT UNIQUE NOT NULL,
    expires_at TIMESTAMP NOT NULL,
    used BOOLEAN DEFAULT 0,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
)
""")

# --- Email verification tokens ---
purchases_cur.execute("""
CREATE TABLE IF NOT EXISTS email_verification_tokens (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    email TEXT NOT NULL,
    token TEXT UNIQUE NOT NULL,
    expires_at TIMESTAMP NOT NULL,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
)
""")

# --- Search activity logging ---
purchases_cur.execute("""
CREATE TABLE IF NOT EXISTS search_logs (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    user_email TEXT,
    search_query TEXT NOT NULL,
    timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    was_blocked BOOLEAN DEFAULT 0
)
""")

# --- Adjudicator page view logging ---
purchases_cur.execute("""
CREATE TABLE IF NOT EXISTS adjudicator_page_views (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    user_email TEXT,
    adjudicator_name TEXT NOT NULL,
    timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP
)
""")

# --- Feedback submissions ---
purchases_cur.execute("""
CREATE TABLE IF NOT EXISTS feedback (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    type TEXT NOT NULL,
    name TEXT,
    email TEXT,
    subject TEXT NOT NULL,
    priority TEXT NOT NULL,
    details TEXT NOT NULL,
    browser TEXT,
    device TEXT,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
)
""")

# --- Safely add columns if they don't exist ---
try:
    purchases_cur.execute("ALTER TABLE purchase_users ADD COLUMN email_verified BOOLEAN DEFAULT 0")
    print("Added 'email_verified' column to purchase_users table.")
except sqlite3.OperationalError:
    pass # Column already exists

try:
    purchases_cur.execute("ALTER TABLE adjudicator_purchases ADD COLUMN stripe_invoice_id TEXT")
    print("Added 'stripe_invoice_id' column to adjudicator_purchases table.")
except sqlite3.OperationalError:
    pass # Column already exists

try:
    purchases_cur.execute("ALTER TABLE search_logs ADD COLUMN was_blocked BOOLEAN DEFAULT 0")
    print("Added 'was_blocked' column to search_logs table.")
except sqlite3.OperationalError:
    pass # Column already exists

# --- Commit all changes ---
purchases_con.commit()

ADMIN_EMAILS = {
    "edwardsheppard5@gmail.com", 
    "ejsheppard@icloud.com", 
    "esheppard@tglaw.com.au"
}



# ---------------- LexiFile DB (separate from Sopal qbcc.db) ----------------
LEXIFILE_DB_PATH = "/var/data/lexifile.db"
lexi_con = sqlite3.connect(LEXIFILE_DB_PATH, check_same_thread=False)
lexi_con.row_factory = sqlite3.Row

# Create users table (LexiFile only)
lexi_con.execute("""
CREATE TABLE IF NOT EXISTS users (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    email TEXT UNIQUE NOT NULL,
    hashed_password TEXT NOT NULL,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
)
""")
lexi_con.commit()

# ---------------- Auth setup ----------------
SECRET_KEY = os.getenv("LEXIFILE_SECRET_KEY", "dev-secret-key")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 43200 # 30 days

# This is the critical fix: pointing to the correct login URL
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/purchase-login")

def get_purchase_user_by_email(email: str):
    """Fetches a user from the adjudicator purchases database by email."""
    cur = purchases_con.execute("SELECT * FROM purchase_users WHERE email = ?", (email,))
    row = cur.fetchone()
    return dict(row) if row else None

def get_current_purchase_user(token: str = Depends(oauth2_scheme)):
    """Dependency to get the currently authenticated user for purchase endpoints."""
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email: str = payload.get("sub")
        if email is None:
            raise HTTPException(status_code=401, detail="Invalid token: no subject")
    except JWTError as e:
        raise HTTPException(status_code=401, detail=f"Invalid token: {e}")
    
    user = get_purchase_user_by_email(email)
    if user is None:
        raise HTTPException(status_code=401, detail="User not found")
    return user

def verify_password(plain_password: str, hashed_password: str) -> bool:
    """Verify password using bcrypt directly"""
    return bcrypt.checkpw(plain_password.encode('utf-8'), hashed_password.encode('utf-8'))

def get_password_hash(password: str) -> str:
    """Hash password using bcrypt directly"""
    password_bytes = password.encode('utf-8')[:72]
    salt = bcrypt.gensalt()
    hashed = bcrypt.hashpw(password_bytes, salt)
    return hashed.decode('utf-8')

def create_access_token(data: dict, expires_delta: timedelta | None = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)



# --- SMTP config (update around line 95) ---
SMTP_FROM_EMAIL = os.getenv("SMTP_FROM_EMAIL", "info@sopal.com.au")
SMTP_HOST = os.getenv("SMTP_HOST", "smtp.gmail.com")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USERNAME = os.getenv("SMTP_USERNAME", "info@sopal.com.au")
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD")


@app.post("/request-password-reset")
async def request_password_reset(email: str = Form(...)):
    """Generates a password reset token and sends email"""
    try:
        # Check if user exists
        user = purchases_con.execute(
            "SELECT email FROM purchase_users WHERE email = ?", (email,)
        ).fetchone()
        
        if not user:
            # Don't reveal if email exists for security
            print(f"Password reset requested for non-existent email: {email}")
            return {"msg": "If that email exists, a reset link has been sent"}
        
        # Generate secure token
        token = secrets.token_urlsafe(32)
        expires_at = datetime.utcnow() + timedelta(hours=1)
        
        # Store token
        purchases_con.execute("""
            INSERT INTO password_reset_tokens (email, token, expires_at)
            VALUES (?, ?, ?)
        """, (email, token, expires_at))
        purchases_con.commit()
        
        print(f"Password reset token created for: {email}")
        
        # Send email with reset link
        reset_link = f"https://sopal.com.au/reset-password?token={token}"
        
        msg = EmailMessage()
        msg["From"] = os.getenv("SMTP_FROM_EMAIL", "info@sopal.com.au")
        msg["To"] = email
        msg["Subject"] = "Password Reset - Sopal"
        msg.set_content(f"""
Hello

You requested a password reset for your Sopal account.

Click the link below to reset your password (valid for 1 hour):
{reset_link}

If you did not request this, please do not click the above link and contact info@sopal.com.au.

Kind regards
Sopal Team
        """)
        
        print(f"Attempting to send password reset email to: {email}")
        
        await aiosmtplib.send(
            msg,
            hostname=os.getenv("SMTP_HOST", "smtp.office365.com"),
            port=int(os.getenv("SMTP_PORT", "587")),
            start_tls=True,
            username=os.getenv("SMTP_USERNAME", "info@sopal.com.au"),
            password=os.getenv("SMTP_PASSWORD"),
        )
        
        print(f"Password reset email sent successfully to: {email}")
        
        return {"msg": "If that email exists, a reset link has been sent"}
        
    except Exception as e:
        print(f"Password reset request error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Failed to process request")






@app.post("/reset-password")
async def reset_password(
    token: str = Form(...),
    new_password: str = Form(...)
):
    """Resets password using valid token"""
    try:
        token_record = purchases_con.execute("""
            SELECT email, expires_at, used 
            FROM password_reset_tokens 
            WHERE token = ?
        """, (token,)).fetchone()
        
        if not token_record:
            raise HTTPException(status_code=400, detail="Invalid or expired reset link")
        
        if token_record['used']:
            raise HTTPException(status_code=400, detail="This reset link has already been used")
        
        expires_at = datetime.fromisoformat(token_record['expires_at'])
        if datetime.utcnow() > expires_at:
            raise HTTPException(status_code=400, detail="This reset link has expired")
        
        hashed_password = get_password_hash(new_password)
        purchases_con.execute("""
            UPDATE purchase_users 
            SET hashed_password = ? 
            WHERE email = ?
        """, (hashed_password, token_record['email']))
        
        purchases_con.execute("""
            UPDATE password_reset_tokens 
            SET used = 1 
            WHERE token = ?
        """, (token,))
        
        purchases_con.commit()
        
        return {"msg": "Password reset successful", "email": token_record['email']}
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Password reset error: {e}")
        raise HTTPException(status_code=500, detail="Failed to reset password")


# Create email verification tokens table
purchases_cur.execute("""
CREATE TABLE IF NOT EXISTS email_verification_tokens (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    email TEXT NOT NULL,
    token TEXT UNIQUE NOT NULL,
    expires_at TIMESTAMP NOT NULL,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
)
""")

# Add email_verified column to users table if it doesn't exist
# This handles existing databases that might already have the column
try:
    purchases_cur.execute("""
    ALTER TABLE purchase_users ADD COLUMN email_verified BOOLEAN DEFAULT 0
    """)
    purchases_con.commit()
    print("Added email_verified column to purchase_users table")
except sqlite3.OperationalError as e:
    # Column already exists, ignore the error
    if "duplicate column name" in str(e).lower():
        print("email_verified column already exists - skipping")
    else:
        print(f"Note: Could not add email_verified column: {e}")



@app.post("/send-verification-email")
async def send_verification_email(current_user: dict = Depends(get_current_purchase_user)):
    """Sends email verification link"""
    try:
        token = secrets.token_urlsafe(32)
        expires_at = datetime.utcnow() + timedelta(hours=24)
        
        purchases_con.execute("""
            INSERT INTO email_verification_tokens (email, token, expires_at)
            VALUES (?, ?, ?)
        """, (current_user['email'], token, expires_at))
        purchases_con.commit()
        
        # CHANGED: Use sopal.com.au domain
        verify_link = f"https://sopal.com.au/verify-email?token={token}"
        
        msg = EmailMessage()
        msg["From"] = SMTP_FROM_EMAIL
        msg["To"] = current_user['email']
        msg["Subject"] = "Verify Your Email - Sopal"
        msg.set_content(f"""
Hello {current_user.get('first_name', '')}

Please verify your email address by clicking the link below:
{verify_link}

This link will expire in 24 hours.

Kind regards
Sopal Team
        """)
        
        await aiosmtplib.send(
            msg,
            hostname=SMTP_HOST,
            port=SMTP_PORT,
            start_tls=True,
            username=SMTP_USERNAME,
            password=SMTP_PASSWORD,
        )
        
        return {"msg": "Verification email sent"}
        
    except Exception as e:
        print(f"Send verification error: {e}")
        raise HTTPException(status_code=500, detail="Failed to send verification email")


@app.get("/verify-email")
async def verify_email(token: str = Query(...)):
    """Verifies email using token and redirects to account page"""
    try:
        token_record = purchases_con.execute("""
            SELECT email, expires_at 
            FROM email_verification_tokens 
            WHERE token = ?
        """, (token,)).fetchone()
        
        if not token_record:
            # Redirect to account page without success message
            return RedirectResponse(url="/account.html?tab=profile", status_code=303)
        
        expires_at = datetime.fromisoformat(token_record['expires_at'])
        if datetime.utcnow() > expires_at:
            # Redirect to account page without success message
            return RedirectResponse(url="/account.html?tab=profile", status_code=303)
        
        purchases_con.execute("""
            UPDATE purchase_users 
            SET email_verified = 1 
            WHERE email = ?
        """, (token_record['email'],))
        
        purchases_con.execute("""
            DELETE FROM email_verification_tokens 
            WHERE token = ?
        """, (token,))
        
        purchases_con.commit()
        
        # Redirect to account page with success indicator
        return RedirectResponse(url="/account.html?tab=profile&verified=success", status_code=303)
        
    except Exception as e:
        print(f"Email verification error: {e}")
        # Redirect to account page on error
        return RedirectResponse(url="/account.html?tab=profile", status_code=303)




# --- CORS Middleware ---
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- In-Memory Databases for LexiFile Demonstration ---
PROJECTS_DB = []
DOCUMENTS_DB = {} 
ARTIFACT_ID_COUNTER = 1


# --- RBA Interest Rate Setup ---
def setup_rba_table():
    con.execute("""
    CREATE TABLE IF NOT EXISTS rba_rates (
        rate_date DATE PRIMARY KEY,
        rate_value REAL NOT NULL
    )
    """)
    con.commit()

def fetch_and_update_rba_rates():
    print("Scheduler: Starting RBA rate update job...")
    try:
        url = "https://www.rba.gov.au/statistics/tables/xls/f01d.xlsx"
        
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36'}
        response = requests.get(url, headers=headers)
        response.raise_for_status()

        excel_data = io.BytesIO(response.content)
        df = pd.read_excel(excel_data, sheet_name='Data', header=1)
        
        date_col = 'Title' 
        rate_col = 'EOD 3-month BABs/NCDs'

        if date_col not in df.columns or rate_col not in df.columns:
            print(f"Scheduler ERROR: Required columns ('{date_col}', '{rate_col}') not found in RBA file.")
            return

        rates_df = df[[date_col, rate_col]].copy()
        rates_df.columns = ['rate_date', 'rate_value']
        
        rates_df['rate_date'] = pd.to_datetime(rates_df['rate_date'], errors='coerce')
        rates_df.dropna(subset=['rate_date', 'rate_value'], inplace=True)
        
        rates_df['rate_date'] = rates_df['rate_date'].dt.date

        cursor = con.cursor()
        for index, row in rates_df.iterrows():
            cursor.execute("""
                INSERT OR IGNORE INTO rba_rates (rate_date, rate_value) VALUES (?, ?)
            """, (row['rate_date'], row['rate_value']))
        con.commit()
        print(f"Scheduler: RBA rate update complete. {len(rates_df)} rows processed.")

    except Exception as e:
        print(f"Scheduler ERROR: Failed to fetch or update RBA rates. Error: {e}")

# Setup and run the scheduler
setup_rba_table()
scheduler = BackgroundScheduler()
scheduler.add_job(fetch_and_update_rba_rates, 'interval', days=1)
scheduler.start()
# Run once on startup as well
fetch_and_update_rba_rates()


# OpenAI
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
con.execute("""
CREATE TABLE IF NOT EXISTS ai_summaries (
    decision_id TEXT PRIMARY KEY,
    summary TEXT
)
""")
con.commit()



# ---------------- ensure FTS ----------------
def ensure_fts():
    try:
        con.execute("DROP TABLE IF EXISTS fts;")
        con.execute("CREATE VIRTUAL TABLE fts USING fts4(full_text, content='docs_fresh', tokenize=unicode61);")
        con.execute("INSERT INTO fts(fts) VALUES('rebuild');")
        con.commit()
    except Exception as e:
        print("Error rebuilding FTS index:", e)

ensure_fts()

# ---------------- meilisearch config ----------------
MEILI_URL = os.getenv("MEILI_URL", "http://127.0.0.1:7700")
MEILI_KEY = os.getenv("MEILI_MASTER_KEY", "")
MEILI_INDEX = "decisions"


# ---------------- helpers ----------------
def normalize_query(q: str) -> str:
    s = q or ""
    for _ in range(2):
        s2 = unquote_plus(s)
        if s2 == s:
            break
        s = s2
    s = unicodedata.normalize("NFKC", s)
    
    replacements = {
        'â€œ': '"', 'â€': '"', 'â€˜': "'", 'â€™': "'", 
        'â€"': '-', 'â€"': '-', 'â€': '-'
    }
    for old, new in replacements.items():
        s = s.replace(old, new)
    
    s = re.sub(r"\s+", " ", s).strip()
    return s

def escape_fts_phrase(s: str) -> str:
    return s.replace('"', '""')

def normalize_default(q: str) -> str:
    if not re.search(r'"|\bNEAR/\d+\b|\bw/\d+\b|\bAND\b|\bOR\b|\bNOT\b|\(|\)', q or "", flags=re.I):
        toks = re.findall(r'\w+\*?', q or "")
        q = ' AND '.join(toks)
    return q

def _fix_unbalanced_quotes(s: str) -> str:
    if s.count('"') % 2 == 1:
        s = s + '"'
    return s

def expand_wildcards(q: str) -> str:
    print(f"Before wildcard expansion: {q}")
    expanded = re.sub(r'(\w+)!', r'\1*', q)
    print(f"After wildcard expansion: {expanded}")
    return expanded

def _parse_near_robust(q: str) -> str | None:
    s = _fix_unbalanced_quotes(q)
    
    if not re.search(r'\b(?:w|near)\s*/\s*\d+\b', s, flags=re.I):
        return None
    
    s = re.sub(r'\b(?:w|near)\s*/\s*(\d+)\b', r'NEAR/\1', s, flags=re.I)

    near_match = re.search(r'NEAR/(\d+)', s, flags=re.I)
    if not near_match:
        return None
    dist = int(near_match.group(1))
    
    parts = re.split(r'\s+NEAR/\d+\s+', s, flags=re.I)
    if len(parts) == 2:
        left_part = parts[0].strip()
        right_part = parts[1].strip()

        def clean_term(term):
            term = term.strip()
            while term.startswith('(') and term.endswith(')'):
                term = term[1:-1].strip()
            if term.startswith('"') and term.endswith('"'):
                term = term[1:-1].strip()
            return term

        left_clean = clean_term(left_part)
        right_clean = clean_term(right_part)
            
        print(f"_parse_near_robust - left: '{left_clean}', dist: {dist}, right: '{right_clean}'")
        return f'"{left_clean}" NEAR/{dist} "{right_clean}"'
    
    return None

def preprocess_sqlite_query(q: str) -> str:
    print(f"preprocess_sqlite_query input: {q}")
    
    q = expand_wildcards(q)
    print(f"After expand_wildcards: {q}")
    
    q = re.sub(r'\band\b', 'AND', q, flags=re.I)
    q = re.sub(r'\bor\b', 'OR', q, flags=re.I)  
    q = re.sub(r'\bnot\b', 'NOT', q, flags=re.I)
    q = re.sub(r'\b(?:w|near)\s*/\s*(\d+)\b', r'NEAR/\1', q, flags=re.I)
    print(f"After operator normalization: {q}")

    near_group_pattern = re.compile(r'(".*?"|\S+)\s+NEAR/(\d+)\s+\((.*?)\)', flags=re.I)
    m = near_group_pattern.search(q)
    
    if m:
        left_term = m.group(1).strip()
        dist = m.group(2)
        right_group = m.group(3).strip()
        
        operator = None
        terms = []
        if ' AND ' in right_group:
            operator = ' AND '
            terms = re.split(r'\s+AND\s+', right_group, flags=re.I)
        elif ' OR ' in right_group:
            operator = ' OR '
            terms = re.split(r'\s+OR\s+', right_group, flags=re.I)
        
        if operator and len(terms) > 1:
            expanded_clauses = []
            for term in terms:
                clean_left = left_term.strip().strip('"')
                clean_right = term.strip().strip('"')
                expanded_clauses.append(f'("{clean_left}" NEAR/{dist} "{clean_right}")')
            
            final_query = operator.join(expanded_clauses)
            print(f"Expanded NEAR/{operator.strip()} query to: {final_query}")
            return final_query

    near_expr = _parse_near_robust(q)
    if near_expr:
        print(f"Found NEAR expression: {near_expr}")
        return near_expr

    if re.search(r'\b(AND|OR|NOT)\b', q, flags=re.I):
        print(f"Contains boolean operators, returning: {q.strip()}")
        return q.strip()

    result = normalize_default(q)
    print(f"normalize_default result: {result}")
    return result

def get_highlight_terms(query: str) -> tuple[list[str], list[str]]:
    phrase_terms = []
    word_terms = []
    
    phrases = re.findall(r'"([^"]+)"', query)
    for phrase in phrases:
        phrase_terms.append(phrase)
    
    query_no_quotes = re.sub(r'"[^"]*"', '', query)
    words = re.findall(r'\b\w+\*?\b', query_no_quotes)
    
    operators = {'AND', 'OR', 'NOT', 'NEAR', 'W'}
    for word in words:
        if (word.upper() not in operators and 
            not word.isdigit() and 
            not re.match(r'\d+', word)):
            word_terms.append(word)
    
    print(f"get_highlight_terms - phrase_terms: {phrase_terms}, word_terms: {word_terms}")
    return phrase_terms, word_terms

_word_re = re.compile(r"\w+", flags=re.UNICODE)

def _tokenize(text: str):
    return _word_re.findall(text.lower())

def _phrase_positions(tokens: list[str], phrase: str) -> list[int]:
    words = _word_re.findall(phrase.lower())
    if not words:
        return []
    if len(words) == 1:
        w = words[0]
        return [i for i, t in enumerate(tokens) if t == w]
    L = len(words)
    out = []
    for i in range(0, len(tokens) - L + 1):
        if tokens[i:i+L] == words:
            out.append(i)
    return out

def _extract_near_components(nq2: str):
    m = re.search(r'"([^"]+)"\s+NEAR/(\d+)\s+"([^"]+)"', nq2, flags=re.I)
    if not m:
        return None
    return m.group(1), int(m.group(2)), m.group(3)

def _filter_rows_for_true_phrase_near(rows, nq2):
    comp = _extract_near_components(nq2)
    if not comp:
        return rows
    left, dist, right = comp
    left_is_phrase  = ' ' in left.strip()
    right_is_phrase = ' ' in right.strip()
    if not (left_is_phrase or right_is_phrase):
        return rows

    filtered = []
    rowids = [r["rowid"] for r in rows]
    if not rowids:
        return rows

    qmarks = ",".join(["?"]*len(rowids))
    text_rows = con.execute(f"SELECT rowid, full_text FROM docs_fresh WHERE rowid IN ({qmarks})", rowids).fetchall()
    text_map = {tr["rowid"]: tr["full_text"] or "" for tr in text_rows}

    for r in rows:
        ft = text_map.get(r["rowid"], "")
        if not ft:
            continue
        toks = _tokenize(ft)
        left_positions  = _phrase_positions(toks, left)
        right_positions = _phrase_positions(toks, right)
        if not left_positions or not right_positions:
            continue

        ok = False
        for li in left_positions:
            for ri in right_positions:
                if abs(ri - li) <= dist:
                    ok = True
                    break
            if ok:
                break
        if ok:
            filtered.append(r)
    return filtered

def highlight_wildcard_matches(text: str, stem: str) -> str:
    pattern = f'\\b({re.escape(stem)}\\w*)'
    return re.sub(pattern, r'<mark>\1</mark>', text, flags=re.I)

# ---------------- routes ----------------
@app.get("/health")
def health():
    return {"ok": True}

@app.get("/get_interest_rate")
def get_interest_rate(start_date_str: str = Query(..., alias="startDate"), end_date_str: str = Query(..., alias="endDate")):
    cursor = con.cursor()
    try:
        start_date = datetime.strptime(start_date_str, "%Y-%m-%d").date()
        end_date = datetime.strptime(end_date_str, "%Y-%m-%d").date()
        
        # Create a temporary table of all dates in the range
        cursor.execute("CREATE TEMP TABLE date_range (day DATE)")
        current_date = start_date
        while current_date <= end_date:
            cursor.execute("INSERT INTO date_range (day) VALUES (?)", (current_date,))
            current_date += timedelta(days=1)

        # For each date in our range, find the most recent RBA rate
        cursor.execute("""
            SELECT 
                dr.day,
                (SELECT rate_value FROM rba_rates WHERE rate_date <= dr.day ORDER BY rate_date DESC LIMIT 1) as rate
            FROM date_range dr
        """)
        
        rows = cursor.fetchall()
        
        if not rows:
             raise HTTPException(status_code=404, detail="No interest rates found for the specified date range.")

        # Convert rows to a list of dictionaries
        daily_rates = [{"date": row["day"], "rate": row["rate"]} for row in rows if row["rate"] is not None]

        if not daily_rates:
             raise HTTPException(status_code=404, detail="No valid interest rates could be mapped to the specified date range.")

        return {"dailyRates": daily_rates}

    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid date format. Please use YYYY-MM-DD.")
    except Exception as e:
        print(f"Error in /get_interest_rate: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")
    finally:
        # Ensure the temporary table is dropped
        cursor.execute("DROP TABLE IF EXISTS date_range")

from fastapi import Depends

from fastapi import Header, Depends # Ensure Header is imported at the top


# In server.py, find the @app.get("/search_fast") route and
# replace the ENTIRE function with this corrected version.
# Make sure you have `from fastapi import Query` at the top of your file.

@app.get("/search_fast")
def search_fast(
    q: str = "",
    limit: int = 20,
    offset: int = 0,
    sort: str = "newest",
    authorization: Optional[str] = Header(None),
    startDate: Optional[str] = Query(None),
    endDate: Optional[str] = Query(None),
    minClaim: Optional[float] = Query(None),
    maxClaim: Optional[float] = Query(None),
    minAdjudicated: Optional[float] = Query(None),
    maxAdjudicated: Optional[float] = Query(None)
):
    # --- START: User and Search Logging ---
    user_email = "Anonymous"
    if authorization and authorization.startswith("Bearer "):
        token = authorization.split(" ")[1]
        try:
            payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
            user_email = payload.get("sub", "Anonymous")
        except JWTError:
            pass  # Token is invalid, treat as Anonymous

    if q:
        purchases_con.execute(
            "INSERT INTO search_logs (user_email, search_query, was_blocked) VALUES (?, ?, 0)",
            (user_email, q)
        )
        purchases_con.commit()
    # --- END: User and Search Logging ---

    q_norm = normalize_query(q)

    # --- START: Hybrid Search Logic ---
    # Determine which search engine to use.
    # Use advanced SQLite FTS if any advanced filters are applied or if the query syntax is complex.
    # Otherwise, use the faster MeiliSearch for simple queries.

    has_advanced_filters = any([
        startDate, endDate, minClaim is not None, maxClaim is not None,
        minAdjudicated is not None, maxAdjudicated is not None
    ])

    nq_expanded = expand_wildcards(q_norm)
    is_complex_query = (
        re.search(r'\b(AND|OR|NOT)\b', q_norm, flags=re.I) or
        re.search(r'\bw/\d+\b', q_norm, flags=re.I) or
        re.search(r'\bNEAR/\d+\b', q_norm, flags=re.I) or
        q_norm.startswith('"') or
        '!' in q_norm or
        '*' in nq_expanded or
        q_norm != nq_expanded
    )
    
    # --- BRANCH 1: SQLite FTS for Advanced/Complex Queries ---
    if has_advanced_filters or is_complex_query:
        try:
            nq2 = preprocess_sqlite_query(q_norm) if q_norm else ""

            where_clauses = []
            sql_params = []
            
            # Always filter out results that are missing critical information
            where_clauses.extend([
                "a.decision_date IS NOT NULL",
                "a.decision_date != ''",
                "a.claimant_name IS NOT NULL",
                "a.claimant_name != ''"
            ])

            if nq2:
                where_clauses.append("fts MATCH ?")
                sql_params.append(nq2)

            # Base join for all queries
            base_join = """
                FROM fts
                JOIN docs_fresh d ON fts.rowid = d.rowid
                LEFT JOIN decision_details a ON d.ejs_id = a.ejs_id
            """

            # Add filter conditions
            if startDate:
                where_clauses.append("a.decision_date >= ?")
                sql_params.append(startDate)
            if endDate:
                where_clauses.append("a.decision_date <= ?")
                sql_params.append(endDate)
            if minClaim is not None:
                where_clauses.append("CAST(a.claimed_amount AS REAL) >= ?")
                sql_params.append(minClaim)
            if maxClaim is not None:
                where_clauses.append("CAST(a.claimed_amount AS REAL) <= ?")
                sql_params.append(maxClaim)
            if minAdjudicated is not None:
                where_clauses.append("CAST(a.adjudicated_amount AS REAL) >= ?")
                sql_params.append(minAdjudicated)
            if maxAdjudicated is not None:
                where_clauses.append("CAST(a.adjudicated_amount AS REAL) <= ?")
                sql_params.append(maxAdjudicated)

            where_sql = f"WHERE {' AND '.join(where_clauses)}" if where_clauses else ""

            # Execute COUNT query
            count_sql = f"SELECT COUNT(DISTINCT fts.rowid) {base_join} {where_sql}"
            total = con.execute(count_sql, tuple(sql_params)).fetchone()[0]

            # Determine ORDER BY clause
            order_clause = "ORDER BY rank"  # Default relevance for text queries
            if not q_norm and sort == "relevance": # If no query, relevance is meaningless
                 sort = "newest"
            
            if sort == "newest":
                order_clause = "ORDER BY a.decision_date DESC"
            elif sort == "oldest":
                order_clause = "ORDER BY a.decision_date ASC"
            elif sort == "atoz":
                order_clause = "ORDER BY a.claimant_name ASC"
            elif sort == "ztoa":
                order_clause = "ORDER BY a.claimant_name DESC"
            elif sort == "claim_high":
                order_clause = "ORDER BY CASE WHEN a.claimed_amount IS NULL OR a.claimed_amount = 'N/A' OR a.claimed_amount = '' THEN -1 ELSE CAST(a.claimed_amount AS REAL) END DESC"
            elif sort == "claim_low":
                order_clause = "ORDER BY CASE WHEN a.claimed_amount IS NULL OR a.claimed_amount = 'N/A' OR a.claimed_amount = '' THEN 9999999999 ELSE CAST(a.claimed_amount AS REAL) END ASC"
            elif sort == "adj_high":
                order_clause = "ORDER BY CASE WHEN a.adjudicated_amount IS NULL OR a.adjudicated_amount = 'N/A' OR a.adjudicated_amount = '' THEN -1 ELSE CAST(a.adjudicated_amount AS REAL) END DESC"
            elif sort == "adj_low":
                order_clause = "ORDER BY CASE WHEN a.adjudicated_amount IS NULL OR a.adjudicated_amount = 'N/A' OR a.adjudicated_amount = '' THEN 9999999999 ELSE CAST(a.adjudicated_amount AS REAL) END ASC"
            
            # Build and execute the main query
            snippet_select = "snippet(fts, 0, '<mark>', '</mark>', ' … ', 100)" if nq2 else "substr(d.full_text, 1, 500) || '...'"
            
            sql = f"""
                SELECT DISTINCT fts.rowid, {snippet_select} AS snippet
                {base_join}
                {where_sql}
                {order_clause}
                LIMIT ? OFFSET ?
            """
            
            final_params = tuple(sql_params) + (limit, offset)
            rows = con.execute(sql, final_params).fetchall()

        except sqlite3.OperationalError as e:
            print("FTS Query error:", e)
            raise HTTPException(status_code=500, detail=f"Search query failed: {e}")

        # Process SQLite results
        items = []
        for r in rows:
            meta = con.execute("""
            SELECT a.claimant_name, a.respondent_name, a.adjudicator_name,
                    a.act_category, d.reference, d.pdf_path, d.ejs_id,
                    a.claimed_amount, a.adjudicated_amount, 
                    a.fee_claimant_proportion, a.fee_respondent_proportion,
                    a.decision_date
            FROM docs_fresh d
            LEFT JOIN decision_details a ON d.ejs_id = a.ejs_id
            WHERE d.rowid = ?
            """, (r["rowid"],)).fetchone()

            d = dict(meta) if meta else {}
            d["id"] = d.get("ejs_id", r["rowid"])
            
            # Add backward-compatible field names for frontend
            d["claimant"] = d.get("claimant_name")
            d["respondent"] = d.get("respondent_name")
            d["adjudicator"] = d.get("adjudicator_name")
            d["decision_date_norm"] = d.get("decision_date")
            d["act"] = d.get("act_category")
            
            snippet_raw = r["snippet"]
            
            # Quick fix: Remove leading 0s and clean up highlighting
            if snippet_raw and q_norm:
                # Remove leading 0s from the snippet
                snippet_raw = re.sub(r'\b0+(\w+)', r'\1', snippet_raw)
                
                # Extract only actual search words (no operators, no numbers)
                search_words = re.findall(r'\b\w+\b', q_norm)
                search_words = [w for w in search_words if w.upper() not in ['AND', 'OR', 'NOT', 'NEAR', 'W'] and not w.isdigit() and len(w) > 1]
                
                # Remove existing marks
                snippet_raw = snippet_raw.replace('<mark>', '').replace('</mark>', '')
                
                # Re-highlight only the actual search terms
                for word in search_words:
                    snippet_raw = re.sub(r'\b(' + re.escape(word) + r')\b', r'<mark>\1</mark>', snippet_raw, flags=re.IGNORECASE)
            
            d["snippet"] = snippet_raw
            items.append(d)


        return {"total": total, "items": items}
        
    # --- BRANCH 2: MeiliSearch for Simple Queries ---
    else:
        payload = {
            "q": q_norm,
            "limit": limit,
            "offset": offset,
            "attributesToRetrieve": [
                "id", "reference", "pdf_path",
                "claimant", "respondent", "adjudicator",  # OLD field names in Meili index
                "date", "act", "content"
            ],
            "attributesToHighlight": ["content"],
            "highlightPreTag": "<mark>",
            "highlightPostTag": "</mark>",
            "attributesToCrop": ["content"],
            "cropLength": 100
        }
        
        if sort == "relevance" and not q_norm:
            sort = "newest"

        if sort == "newest":
            payload["sort"] = ["sortable_date:desc"]
        elif sort == "oldest":
            payload["sort"] = ["sortable_date:asc"]
        elif sort == "atoz":
            payload["sort"] = ["claimant:asc"]
        elif sort == "ztoa":
            payload["sort"] = ["claimant:desc"]
            
        headers = {"Authorization": f"Bearer {MEILI_KEY}"} if MEILI_KEY else {}
        res = requests.post(f"{MEILI_URL}/indexes/{MEILI_INDEX}/search", headers=headers, json=payload)
        data = res.json()
        items = []
        for hit in data.get("hits", []):
            # Fetch additional data from decision_details as it's not in Meili
            extra_data = con.execute("""
                SELECT claimed_amount, adjudicated_amount, 
                       fee_claimant_proportion, fee_respondent_proportion
                FROM decision_details
                WHERE ejs_id = ?
            """, (hit.get("id"),)).fetchone()
            
            snippet = hit.get("_formatted", {}).get("content", "")
            # MeiliSearch still has OLD field names, map them to both old and new
            item = {
                "id": hit.get("id"),
                "reference": hit.get("reference"),
                "pdf_path": hit.get("pdf_path"),
                # Old field names from Meili (for frontend compatibility)
                "claimant": hit.get("claimant"),
                "respondent": hit.get("respondent"),
                "adjudicator": hit.get("adjudicator"),
                # New field names (also map from old)
                "claimant_name": hit.get("claimant"),
                "respondent_name": hit.get("respondent"),
                "adjudicator_name": hit.get("adjudicator"),
                "decision_date_norm": hit.get("date"),
                "decision_date": hit.get("date"),
                "act": hit.get("act"),
                "act_category": hit.get("act"),
                "snippet": snippet
            }
            
            # Add extra fields if available
            if extra_data:
                item.update({
                    "claimed_amount": extra_data["claimed_amount"],
                    "adjudicated_amount": extra_data["adjudicated_amount"],
                    "fee_claimant_proportion": extra_data["fee_claimant_proportion"],
                    "fee_respondent_proportion": extra_data["fee_respondent_proportion"]
                })
            
            items.append(item)
        return {"total": data.get("estimatedTotalHits", 0), "items": items}

    

# ---------- PDF links ----------
GCS_BUCKET = "sopal-bucket"
GCS_PREFIX = "pdfs"

def build_gcs_url(file_name: str) -> str:
    return f"https://storage.googleapis.com/{GCS_BUCKET}/{GCS_PREFIX}/{file_name}"

@app.get("/open")
def open_pdf(p: str, disposition: str = "inline"):
    try:
        file_name = os.path.basename(p)
        return {"url": build_gcs_url(file_name)}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=400)

def get_admin_user(current_user: dict = Depends(get_current_purchase_user)):
    if current_user["email"] not in ADMIN_EMAILS:
        raise HTTPException(status_code=403, detail="Access denied. Admin privileges required.")
    return current_user

# In server.py, add these new endpoints with your other admin routes

@app.get("/admin/full-access-users")
def get_full_access_users(admin: dict = Depends(get_admin_user)):
    """Admin endpoint to get all users with full access."""
    users = purchases_con.execute("SELECT email, added_at FROM full_access_users ORDER BY added_at DESC").fetchall()
    return [dict(row) for row in users]

@app.post("/admin/add-full-access")
def add_full_access_user(email: str = Form(...), admin: dict = Depends(get_admin_user)):
    """Admin endpoint to grant a user full access."""
    try:
        # Check if the user exists in the main user table first
        user_exists = purchases_con.execute("SELECT 1 FROM purchase_users WHERE email = ?", (email,)).fetchone()
        if not user_exists:
            raise HTTPException(status_code=404, detail="User with this email does not have an account.")
        
        purchases_con.execute("INSERT OR IGNORE INTO full_access_users (email) VALUES (?)", (email,))
        purchases_con.commit()
        return {"status": "success", "message": f"{email} granted full access."}
    except sqlite3.IntegrityError:
        return {"status": "success", "message": f"{email} already has full access."}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/admin/remove-full-access")
def remove_full_access_user(email: str = Form(...), admin: dict = Depends(get_admin_user)):
    """Admin endpoint to revoke a user's full access."""
    purchases_con.execute("DELETE FROM full_access_users WHERE email = ?", (email,))
    purchases_con.commit()
    return {"status": "success", "message": f"Full access revoked for {email}."}

@app.get("/admin/all-users")
def get_all_users(admin: dict = Depends(get_admin_user)):
    """Admin endpoint to get all user registration data."""
    users = purchases_con.execute("""
        SELECT id, email, first_name, last_name, firm_name, abn, 
               billing_address, billing_city, billing_state, billing_postcode, 
               employee_size, phone, created_at, last_login 
        FROM purchase_users 
        ORDER BY created_at DESC
    """).fetchall()
    return [dict(row) for row in users]

# Replace your existing get_search_activity function with this one
@app.get("/admin/search-activity")
def get_search_activity(admin: dict = Depends(get_admin_user)):
    """Admin endpoint to get all user search activity."""
    logs = purchases_con.execute("""
        SELECT user_email, search_query, timestamp, was_blocked 
        FROM search_logs 
        ORDER BY timestamp DESC
        LIMIT 1000
    """).fetchall()
    return [dict(row) for row in logs]

# Add this new endpoint with your other admin routes
@app.get("/admin/adjudicator-activity")
def get_adjudicator_activity(admin: dict = Depends(get_admin_user)):
    """Admin endpoint to get adjudicator page view activity."""
    logs = purchases_con.execute("""
        SELECT user_email, adjudicator_name, timestamp
        FROM adjudicator_page_views
        ORDER BY timestamp DESC
        LIMIT 1000
    """).fetchall()
    return [dict(row) for row in logs]

# In server.py, add this new endpoint with your other admin routes
@app.post("/admin/upload-decision")
async def upload_decision(
    admin: dict = Depends(get_admin_user),
    # File
    pdf_file: UploadFile = File(...),
    # IDs
    ejs_id: str = Form(...),
    reference: str = Form(...),
    # Dates
    decision_date: str = Form(...),
    decision_date_norm: str = Form(...),
    # Meta
    act: str = Form(...),
    adjudicator: str = Form(...),
    claimant: str = Form(...),
    respondent: str = Form(...),
    # Financial
    claimed_amount: float = Form(...),
    adjudicated_amount: float = Form(...),
    payment_schedule_amount: Optional[float] = Form(None),
    fee_claimant_proportion: Optional[float] = Form(None),
    fee_respondent_proportion: Optional[float] = Form(None),
    # Classification
    project_type: Optional[str] = Form(None),
    contract_type: Optional[str] = Form(None),
    outcome: Optional[str] = Form(None)
):
    """
    Admin endpoint to upload a new decision, process it, and add to all databases.
    """
    try:
        # 1. Handle the PDF file
        pdf_content = await pdf_file.read()
        pdf_filename = pdf_file.filename

        # 2. Upload PDF to Google Cloud Storage
        storage_client = get_gcs_client()
        gcs_bucket_name = os.getenv("GCS_BUCKET_NAME")
        if not storage_client or not gcs_bucket_name:
            raise HTTPException(status_code=500, detail="GCS not configured on server.")
            
        bucket = storage_client.bucket(gcs_bucket_name)
        blob = bucket.blob(f"pdfs/{pdf_filename}")
        blob.upload_from_string(pdf_content, content_type='application/pdf')
        pdf_path = f"/gcs/{gcs_bucket_name}/pdfs/{pdf_filename}" # Storing path in this format

        # 3. Extract text and page count from the PDF
        full_text = ""
        doc_length_pages = 0
        with io.BytesIO(pdf_content) as f:
            reader = PyPDF2.PdfReader(f)
            doc_length_pages = len(reader.pages)
            for page in reader.pages:
                full_text += page.extract_text() or ""
        
        if not full_text:
            print(f"WARNING: No text could be extracted from {pdf_filename} for ejs_id {ejs_id}")

        # 4. Construct raw_json object
        raw_json_data = {
            "ejs_id": ejs_id,
            "adjudicator_name": adjudicator,
            "claimant_name": claimant,
            "respondent_name": respondent,
            "claimed_amount": claimed_amount,
            "payment_schedule_amount": payment_schedule_amount,
            "adjudicated_amount": adjudicated_amount,
            "jurisdiction_upheld": 0, # Default value
            "fee_claimant_proportion": fee_claimant_proportion,
            "fee_respondent_proportion": fee_respondent_proportion,
            "decision_date": decision_date_norm,
            "keywords": [], # Default value
            "outcome": outcome,
            "sections_referenced": None, # Default value
            "project_type": project_type,
            "contract_type": contract_type,
            "doc_length_pages": doc_length_pages,
            "act_category": act 
        }
        raw_json_string = json.dumps(raw_json_data)

        # 5. Insert data into SQLite database tables
        # Table: docs_fresh
        con.execute(
            "INSERT INTO docs_fresh (ejs_id, reference, pdf_path, full_text) VALUES (?, ?, ?, ?)",
            (ejs_id, reference, pdf_path, full_text)
        )

        # Table: decision_details
        # NOTE: Assuming 'application_no' and 'decision_date' columns exist.
        try:
            con.execute(
                "INSERT INTO decision_details (ejs_id, claimant, respondent, adjudicator, decision_date, decision_date_norm, act, application_no) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                (ejs_id, claimant, respondent, adjudicator, decision_date, decision_date_norm, act, reference)
            )
        except sqlite3.OperationalError as e:
            print(f"Warning: Could not insert into decision_details with extended columns: {e}. Trying basic insert.")
            con.execute(
                "INSERT INTO decision_details (ejs_id, claimant, respondent, adjudicator, decision_date_norm, act) VALUES (?, ?, ?, ?, ?, ?)",
                (ejs_id, claimant, respondent, adjudicator, decision_date_norm, act)
            )


        # Table: decision_details
        con.execute(
            """
            INSERT INTO decision_details (
                ejs_id, adjudicator_name, claimant_name, respondent_name, claimed_amount,
                payment_schedule_amount, adjudicated_amount, fee_claimant_proportion,
                fee_respondent_proportion, decision_date, outcome, project_type, contract_type,
                doc_length_pages, raw_json
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                ejs_id, adjudicator, claimant, respondent, str(claimed_amount),
                str(payment_schedule_amount) if payment_schedule_amount is not None else None,
                str(adjudicated_amount),
                str(fee_claimant_proportion) if fee_claimant_proportion is not None else None,
                str(fee_respondent_proportion) if fee_respondent_proportion is not None else None,
                decision_date_norm, # Using the YYYY-MM-DD date
                outcome, project_type, contract_type,
                doc_length_pages,
                raw_json_string
            )
        )
        con.commit()

        # 6. Update the FTS index for the newly inserted document
        new_doc_row = con.execute("SELECT rowid FROM docs_fresh WHERE ejs_id = ?", (ejs_id,)).fetchone()
        if new_doc_row:
            new_rowid = new_doc_row['rowid']
            con.execute("INSERT INTO fts (rowid, full_text) VALUES (?, ?)", (new_rowid, full_text))
            con.commit()
            print(f"INFO: Successfully updated FTS index for new document ejs_id {ejs_id}")
        else:
            print(f"WARN: Could not find new document {ejs_id} to update FTS index.")

        # 7. Update Meilisearch index
        meili_doc = {
            "id": ejs_id,
            "reference": reference,
            "pdf_path": pdf_path,
            "claimant": claimant,
            "respondent": respondent,
            "adjudicator": adjudicator,
            "date": decision_date_norm,
            "sortable_date": int(time.mktime(datetime.strptime(decision_date_norm, "%Y-%m-%d").timetuple())),
            "act": act,
            "content": full_text
        }
        headers = {"Authorization": f"Bearer {MEILI_KEY}"} if MEILI_KEY else {}
        res = requests.post(f"{MEILI_URL}/indexes/{MEILI_INDEX}/documents", headers=headers, json=[meili_doc])
        res.raise_for_status() # Raise an exception if Meilisearch update fails

        return {"status": "success", "message": f"Decision '{reference}' uploaded and indexed successfully."}

    except Exception as e:
        # Rollback DB changes if anything fails
        con.rollback()
        print(f"ERROR in /admin/upload-decision: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"An error occurred during upload: {e}")

# In server (3).py, REPLACE the entire /admin/rebuild-indexes function with this:

        # In server (3).py, REPLACE the entire /admin/rebuild-indexes function with this:

@app.get("/admin/rebuild-indexes")
async def rebuild_indexes():
    """
    Forces a rebuild of the local FTS index and resynchronizes all documents with Meilisearch
    in a memory-efficient way.
    """
    try:
        # 1. Rebuild the local SQLite FTS index
        print("INFO: Starting FTS index rebuild...")
        ensure_fts()
        print("INFO: FTS index rebuild complete.")

        # 2. Resynchronize with Meilisearch
        print("INFO: Starting Meilisearch synchronization...")
        headers = {"Authorization": f"Bearer {MEILI_KEY}"} if MEILI_KEY else {}
        
        # --- FIX: Clear the index at the very beginning ---
        print("INFO: Deleting all existing documents from Meilisearch...")
        try:
            res = requests.delete(f"{MEILI_URL}/indexes/{MEILI_INDEX}/documents", headers=headers)
            res.raise_for_status()
            print("INFO: Meilisearch index cleared.")
        except Exception as e:
            print(f"WARN: Could not clear Meilisearch index (it might be empty already): {e}")


        # --- MEMORY FIX: Fetch only IDs first, then process in batches ---
        print("INFO: Fetching all document IDs...")
        all_doc_ids = con.execute("SELECT ejs_id FROM docs_fresh WHERE ejs_id IS NOT NULL").fetchall()
        
        if not all_doc_ids:
            return {"status": "warning", "message": "No documents found in the database to sync."}
            
        total_docs = len(all_doc_ids)
        print(f"INFO: Found {total_docs} document IDs to process.")

        batch_size = 200 # Process 200 docs at a time to keep memory low
        
        # Helper function to safely convert values to float
        def to_float(value):
            try:
                return float(value)
            except (ValueError, TypeError):
                return 0.0

        for i in range(0, total_docs, batch_size):
            # Get the list of ejs_ids for the current batch
            batch_ids_tuples = all_doc_ids[i:i + batch_size]
            batch_ids = [row['ejs_id'] for row in batch_ids_tuples]
            placeholders = ",".join("?" * len(batch_ids))

            print(f"INFO: Processing batch {i//batch_size + 1} / {-(total_docs // -batch_size)} (docs {i+1} to {min(i+batch_size, total_docs)})...")

            # Run the full query for *only* this batch of IDs
            docs_to_sync = con.execute(f"""
                SELECT 
                    d.ejs_id, d.reference, d.pdf_path, d.full_text,
                    m.claimant_name, m.respondent_name, m.adjudicator_name,
                    m.decision_date, m.claimed_amount, m.adjudicated_amount,
                    m.fee_claimant_proportion, m.fee_respondent_proportion
                FROM docs_fresh d
                LEFT JOIN decision_details m ON d.ejs_id = m.ejs_id
                WHERE d.ejs_id IN ({placeholders})
            """, batch_ids).fetchall()

            meili_docs = []
            for doc in docs_to_sync:
                try:
                    sortable_date = 0
                    if doc["decision_date"]:
                        sortable_date = int(time.mktime(datetime.strptime(doc["decision_date"], "%Y-%m-%d").timetuple()))

                    meili_doc = {
                        "id": doc["ejs_id"],
                        "reference": doc["reference"],
                        "pdf_path": doc["pdf_path"],
                        "claimant": doc["claimant_name"],
                        "respondent": doc["respondent_name"],
                        "adjudicator": doc["adjudicator_name"],
                        "date": doc["decision_date"],
                        "sortable_date": sortable_date,
                        "act": "BIF Act",
                        "content": doc["full_text"] or "",
                        "claimed_amount": to_float(doc["claimed_amount"]),
                        "adjudicated_amount": to_float(doc["adjudicated_amount"]),
                        "fee_claimant_proportion": to_float(doc["fee_claimant_proportion"]),
                        "fee_respondent_proportion": to_float(doc["fee_respondent_proportion"])
                    }
                    meili_docs.append(meili_doc)
                except (ValueError, TypeError) as e:
                    print(f"WARN: Skipping document {doc['ejs_id']} due to invalid date '{doc['decision_date']}': {e}")
                except Exception as e:
                    print(f"WARN: Skipping document {doc['ejs_id']} due to processing error: {e}")

            # Send this batch to Meilisearch
            if meili_docs:
                res = requests.post(f"{MEILI_URL}/indexes/{MEILI_INDEX}/documents", headers=headers, json=meili_docs)
                res.raise_for_status()
                print(f"INFO: Sent batch {i//batch_size + 1} to Meilisearch.")
            else:
                print(f"INFO: Batch {i//batch_size + 1} had no documents to send.")

        print(f"INFO: Meilisearch synchronization complete. {total_docs} documents processed.")
        
        return {"status": "success", "message": f"Successfully rebuilt SQLite FTS index and synchronized all {total_docs} documents with Meilisearch."}

    except Exception as e:
        print(f"ERROR in /admin/rebuild-indexes: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"An error occurred during index rebuild: {e}")
        




# In server.py, REPLACE the entire create_meilisearch_backup function with this one

@app.post("/admin/create-meilisearch-backup")
def create_meilisearch_backup(admin: dict = Depends(get_admin_user)):
    """
    Triggers a Meilisearch snapshot, waits for it, finds the new file on disk,
    and uploads it to the GCS bucket.
    """
    SNAPSHOT_DIR = "/meili_data/snapshots" # The shared disk path

    try:
        # Step 1: Tell Meilisearch to start creating a snapshot
        headers = {"Authorization": f"Bearer {MEILI_KEY}"} if MEILI_KEY else {}
        start_res = requests.post(f"{MEILI_URL}/snapshots", headers=headers)
        start_res.raise_for_status()
        task = start_res.json()
        task_uid = task.get('taskUid')
        if not task_uid:
            raise Exception("Meilisearch did not return a valid task UID to track.")

        # Step 2: Poll Meilisearch until the snapshot task is complete
        print(f"INFO: Waiting for Meilisearch snapshot task {task_uid} to complete...")
        while True:
            time.sleep(2) # Wait 2 seconds between checks
            task_res = requests.get(f"{MEILI_URL}/tasks/{task_uid}", headers=headers)
            task_res.raise_for_status()
            task_status = task_res.json()

            if task_status.get('status') == 'succeeded':
                print(f"INFO: Snapshot task {task_uid} succeeded.")
                break # Exit the loop on success

            if task_status.get('status') in ['failed', 'canceled']:
                error_details = task_status.get('error', {})
                raise Exception(f"Meilisearch snapshot failed: {error_details.get('message', 'No error details provided.')}")
        
        # --- START: NEW, MORE ROBUST LOGIC ---
        # Step 3: Find the newest snapshot file directly from the shared disk
        print(f"INFO: Searching for the newest snapshot file in '{SNAPSHOT_DIR}'...")
        
        # Ensure the directory exists before listing
        if not os.path.isdir(SNAPSHOT_DIR):
             raise Exception(f"Snapshot directory '{SNAPSHOT_DIR}' not found. Check your Meilisearch start command and disk mount.")

        snapshot_files = [f for f in os.listdir(SNAPSHOT_DIR) if f.endswith('.snapshot')]
        if not snapshot_files:
            raise Exception("Snapshot task succeeded, but no snapshot files were found in the directory.")
            
        # Find the most recently created/modified snapshot file
        latest_snapshot_filename = max(snapshot_files, key=lambda f: os.path.getmtime(os.path.join(SNAPSHOT_DIR, f)))
        local_snapshot_path = os.path.join(SNAPSHOT_DIR, latest_snapshot_filename)
        print(f"INFO: Found latest snapshot file: {latest_snapshot_filename}")
        # --- END: NEW LOGIC ---

        # Step 4: Upload the snapshot from the shared disk to GCS
        print(f"INFO: Uploading '{latest_snapshot_filename}' to GCS...")
        storage_client = get_gcs_client()
        gcs_bucket_name = os.getenv("GCS_BUCKET_NAME")
        if not storage_client or not gcs_bucket_name:
            raise HTTPException(status_code=500, detail="GCS not configured on the server.")

        bucket = storage_client.bucket(gcs_bucket_name)
        gcs_path = f"meili_backups/{latest_snapshot_filename}"
        blob = bucket.blob(gcs_path)
        
        blob.upload_from_filename(local_snapshot_path)
        print(f"INFO: Successfully uploaded to GCS at '{gcs_path}'.")

        # Step 5: Clean up the local snapshot to save disk space
        os.remove(local_snapshot_path)
        print(f"INFO: Cleaned up local file: {local_snapshot_path}")

        return {"status": "success", "message": f"Backup '{latest_snapshot_filename}' successfully uploaded to GCS."}

    except Exception as e:
        print(f"ERROR in /admin/create-meilisearch-backup: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
    
# In server.py, add this with your other admin endpoints

@app.post("/admin/save-database")
def save_database_to_gcs(admin: dict = Depends(get_admin_user)):
    """
    Uploads the current local qbcc.db file from /tmp back to GCS, overwriting the old one.
    """
    try:
        # CRITICAL FIX: Force a checkpoint to merge the WAL file into the main DB file
        print("INFO: Checkpointing database to ensure all changes are written...")
        con.execute("PRAGMA wal_checkpoint(FULL);")
        print("INFO: Checkpoint complete.")

        # DB_PATH is defined at the top of your server.py file as "/tmp/qbcc.db"
        if not os.path.exists(DB_PATH):
            raise HTTPException(status_code=404, detail="Local database file not found.")

        storage_client = get_gcs_client()
        gcs_bucket_name = os.getenv("GCS_BUCKET_NAME")
        gcs_db_object_name = os.getenv("GCS_DB_OBJECT_NAME", "qbcc.db") # The filename in the bucket

        if not storage_client or not gcs_bucket_name:
            raise HTTPException(status_code=500, detail="GCS not configured on the server.")

        print(f"INFO: Uploading local database '{DB_PATH}' to GCS bucket '{gcs_bucket_name}' as '{gcs_db_object_name}'...")

        bucket = storage_client.bucket(gcs_bucket_name)
        blob = bucket.blob(gcs_db_object_name)
        blob.upload_from_filename(DB_PATH)

        print("INFO: Database upload successful.")
        return {"status": "success", "message": "Database successfully saved to Google Cloud Storage."}

    except Exception as e:
        print(f"ERROR in /admin/save-database: {e}")
        raise HTTPException(status_code=500, detail=str(e))

    
# ---------- AI Summarise ----------
@app.get("/summarise/{decision_id}")
def summarise(decision_id: str = Path(...)):
    try:
        row = con.execute("SELECT summary FROM ai_summaries WHERE decision_id = ?", (decision_id,)).fetchone()
        if row:
            return {"id": decision_id, "summary": row["summary"]}

        r = con.execute("SELECT full_text FROM docs_fresh WHERE ejs_id = ?", (decision_id,)).fetchone()
        if not r or not r["full_text"]:
            raise HTTPException(status_code=404, detail="Decision not found")

        text = r["full_text"][:300000]

        resp = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are SopalAI, a legal assistant specialising in construction law. "
                        "Summarise adjudication decisions under Queensland's Security of Payment legislation "
                        "in clear, plain English. "
                        "Respond as if the user is asking you about the decision directly. "
                        "Do not say 'the adjudication decision you provided' or similar. "
                        "Structure the summary into bullet points or short sections with HTML-friendly formatting "
                        "(<strong> for headings, <ul>/<li> for lists) with compact spacing (no double/triple lines)."
                        "Structure the summary using only HTML tags (<strong>, <ul>, <li>, <p>), "
                        "not Markdown (#, ##, ###)."
                    )
                },
                {
                    "role": "user",
                    "content": f"""
Summarise this adjudication decision in 5–7 bullet points, covering:
- The parties and the works
- Payment claim amount and payment schedule response
- Any jurisdictional challenges
- The factual disputes and evidence (provide at least 5 sentences of detail on this point, be thorough)
- The adjudicator's reasoning
- The final outcome, amount awarded, and fee split
Don't actually title it "Summary" or the like please, go straight into bullet points

Decision text:
{text}
"""
                }
            ],
            max_tokens=1024,
        )

        summary = resp.choices[0].message.content.strip()
        con.execute("INSERT OR REPLACE INTO ai_summaries(decision_id, summary) VALUES (?, ?)", (decision_id, summary))
        con.commit()
        return {"id": decision_id, "summary": summary}
    except Exception as e:
        print("ERROR in /summarise:", e)
        return JSONResponse({"error": str(e)}, status_code=500)

@app.post("/log-blocked-search")
async def log_blocked_search(q: str = Form(...), authorization: Optional[str] = Header(None)):
    user_email = "Anonymous"
    if authorization and authorization.startswith("Bearer "):
        token = authorization.split(" ")[1]
        try:
            payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
            user_email = payload.get("sub", "Anonymous")
        except JWTError:
            pass
    
    if q:
        purchases_con.execute(
            "INSERT INTO search_logs (user_email, search_query, was_blocked) VALUES (?, ?, 1)",
            (user_email, q)
        )
        purchases_con.commit()
    return {"status": "logged"}

@app.get("/api/adjudicators")
def get_adjudicators():
    """Get list of all adjudicators with their statistics from decision_details"""
    try:
        # First get all decisions per adjudicator to calculate median
        query_all = """
        SELECT 
            adjudicator_name,
            claimed_amount,
            adjudicated_amount,
            fee_claimant_proportion,
            fee_respondent_proportion
        FROM decision_details
        WHERE adjudicator_name IS NOT NULL 
        AND adjudicator_name != ''
        AND TRIM(adjudicator_name) != ''
        """
        
        all_rows = con.execute(query_all).fetchall()
        
        # Group by adjudicator and calculate individual award rates and fee proportions
        adjudicator_rates = {}
        adjudicator_fees = {}
        adjudicator_zero_awards = {}
        
        for row in all_rows:
            name = row["adjudicator_name"]
            
            # Safely convert to float, skipping 'N/A' and other invalid values
            try:
                claimed = float(row["claimed_amount"]) if row["claimed_amount"] not in ('N/A', '', None) else None
                adjudicated = float(row["adjudicated_amount"]) if row["adjudicated_amount"] not in ('N/A', '', None) else None
            except (ValueError, TypeError):
                continue
            
      
            # Count $0 awards - include ALL decisions where awarded = $0
            if adjudicated is not None and adjudicated == 0:
                if name not in adjudicator_zero_awards:
                    adjudicator_zero_awards[name] = 0
                adjudicator_zero_awards[name] += 1
            
            # Include ALL decisions where we have valid numbers, including $0 awards
            if claimed is not None and adjudicated is not None and claimed > 0:
                rate = min((adjudicated / claimed) * 100, 100.0)
                if name not in adjudicator_rates:
                    adjudicator_rates[name] = []
                adjudicator_rates[name].append(rate)
            
            # Handle fee proportions
            try:
                claimant_fee = float(row["fee_claimant_proportion"]) if row["fee_claimant_proportion"] not in ('N/A', '', None) else None
                respondent_fee = float(row["fee_respondent_proportion"]) if row["fee_respondent_proportion"] not in ('N/A', '', None) else None
                
                if claimant_fee is not None:
                    if name not in adjudicator_fees:
                        adjudicator_fees[name] = {'claimant': [], 'respondent': []}
                    adjudicator_fees[name]['claimant'].append(claimant_fee)
                    if respondent_fee is None or abs((claimant_fee + respondent_fee) - 100) > 0.1:
                        respondent_fee = 100 - claimant_fee
                    adjudicator_fees[name]['respondent'].append(respondent_fee)
            except (ValueError, TypeError):
                continue
        
        # Now get aggregated stats
        query = """
        SELECT 
            adjudicator_name,
            COUNT(*) as total_decisions,
            SUM(CASE 
                WHEN claimed_amount NOT IN ('N/A', '') AND claimed_amount IS NOT NULL 
                THEN CAST(claimed_amount AS REAL) 
                ELSE 0 
            END) as total_claimed,
            SUM(CASE 
                WHEN adjudicated_amount NOT IN ('N/A', '') AND adjudicated_amount IS NOT NULL 
                THEN CAST(adjudicated_amount AS REAL) 
                ELSE 0 
            END) as total_adjudicated
        FROM decision_details
        WHERE adjudicator_name IS NOT NULL 
        AND adjudicator_name != ''
        AND TRIM(adjudicator_name) != ''
        GROUP BY adjudicator_name
        HAVING COUNT(*) >= 0
        ORDER BY total_decisions DESC
        """
        
        rows = con.execute(query).fetchall()
        
        adjudicators = []
        for row in rows:
            name = row["adjudicator_name"]
            total_claimed = float(row["total_claimed"]) if row["total_claimed"] else 0
            total_adjudicated = float(row["total_adjudicated"]) if row["total_adjudicated"] else 0
            
            # Calculate average (mean) award rate from individual decisions
            avg_award_rate = 0
            if name in adjudicator_rates and adjudicator_rates[name]:
                avg_award_rate = sum(adjudicator_rates[name]) / len(adjudicator_rates[name])
            
            # Calculate average fee proportions
            avg_claimant_fee = 0
            avg_respondent_fee = 0
            if name in adjudicator_fees:
                if adjudicator_fees[name]['claimant']:
                    avg_claimant_fee = sum(adjudicator_fees[name]['claimant']) / len(adjudicator_fees[name]['claimant'])
                if adjudicator_fees[name]['respondent']:
                    avg_respondent_fee = sum(adjudicator_fees[name]['respondent']) / len(adjudicator_fees[name]['respondent'])
            
            adjudicator = {
                "id": name.replace(" ", "_").lower(),
                "name": name,
                "totalDecisions": row["total_decisions"],
                "totalClaimAmount": total_claimed,
                "totalAwardedAmount": total_adjudicated,
                "avgAwardRate": avg_award_rate,
                "avgClaimantFeeProportion": avg_claimant_fee,
                "avgRespondentFeeProportion": avg_respondent_fee,
                "zeroAwardCount": adjudicator_zero_awards.get(name, 0)
            }
            adjudicators.append(adjudicator)
        
        return adjudicators
        
    except Exception as e:
        print(f"Error in /api/adjudicators: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)


@app.get("/api/adjudicator/{adjudicator_name}")
def get_adjudicator_decisions(adjudicator_name: str = Path(...), authorization: Optional[str] = Header(None)):
    """Get all decisions for a specific adjudicator from decision_details"""
    # --- START: FINAL PAGE VIEW LOGGING ---
    user_email = "Anonymous"
    if authorization and authorization.startswith("Bearer "):
        token = authorization.split(" ")[1]
        try:
            payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
            user_email = payload.get("sub", "Anonymous")
        except JWTError:
            pass
    
    decoded_name_for_log = unquote_plus(adjudicator_name)
    purchases_con.execute(
        "INSERT INTO adjudicator_page_views (user_email, adjudicator_name) VALUES (?, ?)",
        (user_email, decoded_name_for_log)
    )
    purchases_con.commit()
    # --- END: FINAL PAGE VIEW LOGGING ---
    
    try:
        decoded_name = unquote_plus(adjudicator_name)
        
        query = """
        SELECT 
            a.ejs_id,
            a.adjudicator_name,
            a.claimant_name,
            a.respondent_name,
            a.claimed_amount,
            a.payment_schedule_amount,
            a.adjudicated_amount,
            a.decision_date,
            a.outcome,
            a.project_type,
            a.contract_type,
            a.fee_claimant_proportion,
            a.fee_respondent_proportion,
            d.reference,
            d.pdf_path
        FROM decision_details a
        LEFT JOIN docs_fresh d ON a.ejs_id = d.ejs_id
        WHERE LOWER(TRIM(a.adjudicator_name)) = LOWER(TRIM(?))
        ORDER BY a.decision_date DESC
        """
        
        rows = con.execute(query, (decoded_name,)).fetchall()
        
        decisions = []
        for row in rows:
            claimed = 0
            try:
                claimed = float(row["claimed_amount"]) if row["claimed_amount"] not in ('N/A', '', None) else 0
            except (ValueError, TypeError):
                claimed = 0
                
            adjudicated = 0
            try:
                adjudicated = float(row["adjudicated_amount"]) if row["adjudicated_amount"] not in ('N/A', '', None) else 0
            except (ValueError, TypeError):
                adjudicated = 0
            
            # Parse fee proportions
            claimant_fee = 0
            respondent_fee = 0
            try:
                claimant_fee = float(row["fee_claimant_proportion"]) if row["fee_claimant_proportion"] not in ('N/A', '', None) else 0
                respondent_fee = float(row["fee_respondent_proportion"]) if row["fee_respondent_proportion"] not in ('N/A', '', None) else 0
            except (ValueError, TypeError):
                pass
            
            decision = {
                "id": row["ejs_id"],
                "title": f"{row['claimant_name'] or 'Unknown'} v {row['respondent_name'] or 'Unknown'}",
                "reference": row["reference"],
                "date": row["decision_date"],
                "claimant": row["claimant_name"],
                "respondent": row["respondent_name"],
                "claimAmount": claimed,
                "awardedAmount": adjudicated,
                "claimantFeeProportion": claimant_fee,
                "respondentFeeProportion": respondent_fee,
                "outcome": row["outcome"],
                "projectType": row["project_type"],
                "contractType": row["contract_type"],
                "pdfPath": row["pdf_path"]
            }
            decisions.append(decision)
        
        return decisions
        
    except Exception as e:
        print(f"Error in /api/adjudicator/{adjudicator_name}: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)




@app.post("/ask-rag")
def ask_rag(query: str = Form(...)):
    """Natural language search across all adjudication decisions using RAG"""
    try:
        if not chroma_collection:
            raise HTTPException(status_code=503, detail="RAG search not available")
        
        # Query the vector database
        results = chroma_collection.query(
            query_texts=[query],
            n_results=5
        )
        
        if not results['documents'][0]:
            return {"answer": "I couldn't find any relevant information in the adjudication decisions for your query."}
        
        # Combine retrieved chunks as context
        context_chunks = results['documents'][0]
        metadata_list = results['metadatas'][0]
        
        context = "\n\n---\n\n".join([
            f"[Decision {meta.get('decision_id', 'Unknown')}]\n{doc}"
            for doc, meta in zip(context_chunks, metadata_list)
        ])
        
        # Send to GPT for answer generation
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are SopalAI, an expert assistant for Queensland adjudication decisions. "
                        "Answer questions based on the provided context from adjudication decisions. "
                        "Cite specific decisions when possible. Be concise and accurate."
                    )
                },
                {
                    "role": "user",
                    "content": f"Context from decisions:\n\n{context}\n\nQuestion: {query}"
                }
            ],
            max_tokens=800
        )
        
        answer = resp.choices[0].message.content.strip()
        
        # Return answer with source references
        sources = list(set([meta.get('decision_id', 'Unknown') for meta in metadata_list]))
        
        return {
            "answer": answer,
            "sources": sources[:5]
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error in /ask-rag: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)

# Add this helper endpoint to debug the database schema
@app.get("/debug/schema")
def debug_schema():
    """Debug endpoint to check database schema"""
    try:
        # Get decision_details table schema
        meta_columns = con.execute("PRAGMA table_info(decision_details)").fetchall()
        
        # Get a sample row to see what data looks like
        sample_row = con.execute("SELECT * FROM decision_details LIMIT 1").fetchone()
        
        return {
            "decision_details_columns": [{"name": col[1], "type": col[2]} for col in meta_columns],
            "sample_data": dict(sample_row) if sample_row else None
        }
    except Exception as e:
        return {"error": str(e)}

# ---------- FEEDBACK FORM ----------
@app.post("/send-feedback")
async def send_feedback(
    type: str = Form(...),
    name: str = Form(""),
    email: str = Form(""),
    subject: str = Form(...),
    priority: str = Form(...),
    details: str = Form(...),
    browser: str = Form(""),
    device: str = Form("")
):
    # Save feedback to database as backup
    try:
        purchases_con.execute("""
            INSERT INTO feedback (type, name, email, subject, priority, details, browser, device, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, datetime('now'))
        """, (type, name, email, subject, priority, details, browser, device))
        purchases_con.commit()
    except Exception as db_err:
        print(f"Warning: Could not save feedback to database: {db_err}")

    # Try to send email notification
    smtp_password = os.getenv("SMTP_PASSWORD")
    if not smtp_password:
        # If no SMTP password configured, just save to DB and return success
        print("Warning: SMTP_PASSWORD not configured, feedback saved to database only")
        return {"ok": True, "message": "Feedback received and saved"}

    msg = EmailMessage()
    msg["From"] = "sopal.aus@gmail.com"
    msg["To"] = "sopal.aus@gmail.com"
    msg["Subject"] = f"[{type.upper()}] {subject}"
    body = f"""
Feedback Type: {type}
Name: {name}
Email: {email}
Priority: {priority}
Browser: {browser}
Device: {device}

Details:
{details}
"""
    msg.set_content(body)

    try:
        await aiosmtplib.send(
            msg,
            hostname="smtp.gmail.com",
            port=587,
            start_tls=True,
            username="sopal.aus@gmail.com",
            password=smtp_password,
        )
        return {"ok": True, "message": "Feedback sent successfully"}
    except Exception as e:
        # Email failed but feedback is already saved in DB
        print(f"SMTP error (feedback saved to DB): {e}")
        return {"ok": True, "message": "Feedback received and saved"}

# ---------- AI Ask ----------
@app.post("/ask/{decision_id}")
def ask_ai(decision_id: str = Path(...), question: str = Form(...), history: str = Form("[]")):
    try:
        row = con.execute("SELECT full_text FROM docs_fresh WHERE ejs_id = ?", (decision_id,)).fetchone()
        if not row or not row["full_text"]:
            raise HTTPException(status_code=404, detail="Decision not found")

        text = row["full_text"][:300000] # Increased context for chat

        try:
            chat_history = json.loads(history)
            if not isinstance(chat_history, list):
                chat_history = []
        except json.JSONDecodeError:
            chat_history = []

        system_prompt = {
            "role": "system",
            "content": (
                "You are SopalAI, assisting users with adjudication decisions under the BIF Act or BCIPA. "
                "Respond directly to the user's question as if they asked you about the decision. "
                "Do not say 'the adjudication decision you provided' or 'the text you gave me'. "
                "Write in clear, plain English with headings and bullet points formatted in HTML."
                "Structure the summary using only HTML tags (<strong>, <ul>, <li>, <p>), "
                "not Markdown (#, ##, ###). The user has already seen your initial summary, "
                "so your follow-up answers should be concise and directly address their question."
            )
        }
        
        context_prompt = {"role": "user", "content": f"Here is the context from the decision document. Answer questions based on this. CONTEXT:\n{text}"}

        messages = [system_prompt, context_prompt] + chat_history + [{"role": "user", "content": question}]

        resp = client.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            max_tokens=1024
        )

        answer = resp.choices[0].message.content.strip()
        return {"answer": answer}
    except Exception as e:
        print("ERROR in /ask:", e)
        return JSONResponse({"error": str(e)}, status_code=500)


# ---------- Compliance Checker AI Endpoint ----------

def extract_text_from_file(file: io.BytesIO, filename: str) -> str:
    """Extracts text from PDF, DOCX, or TXT files."""
    text = ""
    try:
        if filename.lower().endswith('.pdf'):
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() or ""
        elif filename.lower().endswith('.docx'):
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif filename.lower().endswith('.txt'):
            text = file.read().decode('utf-8')
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {filename}")
    except Exception as e:
        print(f"Error extracting text from {filename}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to process file: {filename}")
    return text

def get_system_prompt(doc_type: str) -> str:
    """Returns the expert system prompt for the AI based on document type."""
    if doc_type == 'pc':
        return """
You are an expert construction lawyer in Queensland, specialising in the Building Industry Fairness (Security of Payment) Act 2017 (BIF Act). Your tone must be cautious, thorough, and professional, always erring on the side of highlighting potential risks.

Your task is to analyse the provided text from a document purporting to be a payment claim. Critically assess it for compliance against the BIF Act and key principles from relevant case law. For each check, provide a status ('pass', 'warning', or 'fail'), a short title, and a detailed feedback paragraph explaining your reasoning, citing legislative sections and case law where appropriate.

Your entire response must be a single JSON object with a key "checks" which is an array of check objects. Do not include any text outside of this JSON object.

The check objects must cover the following points in detail:

1.  **Claimed Amount Stated (section 68(1)(b))**: Identify if a clear, specific monetary amount is claimed. A 'pass' requires a single, unambiguous figure. A 'fail' results from ambiguity, multiple conflicting amounts, or no stated amount.
2.  **Identification of Works (section 68(1)(a))**: Assess if the work, related goods, and services are identified with sufficient detail. The standard is not overly onerous, but it must be reasonably clear what the claim is for (see *Trask Development Pty Ltd v. Moreton Bay Regional Council*). A generic description like 'works on site' may be a 'fail'. A project name and general description is usually a 'pass'.
3.  **Request for Payment (section 68(1)(c))**: Check for an explicit request for payment. The Act states a document bearing the word ‘invoice’ is sufficient (section 68(3)). Phrases like 'This is a payment claim made under the BIF Act' are strong indicators. A document merely stating an amount owing without a clear demand for payment may 'fail' (e.g., *KDV Sport Pty Ltd v. Muggeridge Constructions Pty Ltd*).
4.  **Reference Date Validity (section 67)**: This is a critical jurisdictional requirement. 'Fail' the document if no reference date is stated or can be clearly identified. If a date is present, give it a 'pass' but include a comment explaining that the AI cannot verify if this date is valid under the contract or if a claim for this date has already been made. Explain that an invalid reference date is a fatal flaw for the claim.
5.  **Timeliness of Claim (section 75)**: This cannot be verified from the document alone, so always assign a 'warning'. The feedback must explain the critical timeframes: 6 months after the work was last carried out for a progress claim (section 75(2)), and the longer of the periods in section 75(3) for a final claim. Emphasise that failure to comply is a complete bar to proceeding.
6.  **Correct Service**: This also cannot be verified from the document text, so always assign a 'warning'. The feedback must stress the importance of serving the claim on the correct party at the address for notices stipulated in the contract, as improper service can invalidate the entire claim.
"""
    elif doc_type == 'ps':
        return """
You are an expert construction lawyer in Queensland, specialising in the Building Industry Fairness (Security of Payment) Act 2017 (BIF Act). Your tone must be cautious, thorough, and professional, always erring on the side of highlighting potential risks for the respondent.

Your task is to analyse the provided text from a document purporting to be a payment schedule. Critically assess it for compliance against the BIF Act and key principles from relevant case law. For each check, provide a status ('pass', 'warning', or 'fail'), a short title, and a detailed feedback paragraph explaining your reasoning, citing legislative sections and case law where appropriate.

Your entire response must be a single JSON object with a key "checks" which is an array of check objects. Do not include any text outside of this JSON object.

The check objects must cover the following points in detail:

1.  **Identifies Payment Claim (section 69(a))**: Check if the schedule clearly and unambiguously identifies the payment claim it is responding to (e.g., by date, claim number, or project reference). A 'fail' here could render the entire schedule invalid.
2.  **Scheduled Amount Stated (section 69(b))**: Check if it states the amount of the payment, if any, that the respondent proposes to make. This can be zero, but a figure must be stated. Failure to do so is a 'fail'.
3.  **Reasons for Withholding (section 69(c))**: This is the most critical compliance point. The reasons for withholding payment must be articulated with sufficient particularity. A 'fail' is warranted for vague, generic reasons like 'defective works' or 'incomplete works' without further detail. The reasons must be comprehensive enough for the claimant to understand the case they have to meet at adjudication (*John Holland Pty Ltd v. TAC Pacific Pty Ltd*). Emphasise that the respondent is bound by the reasons in their schedule and cannot introduce new reasons later.
4.  **Timeliness of Schedule (section 76)**: This cannot be verified from the document alone, so always assign a 'warning'. The feedback must explain the strict timeframe for service: 15 business days after receiving the payment claim or a shorter period if specified in the contract) and state that failure to provide a compliant schedule on time makes the respondent liable for the full claimed amount.
5.  **Correct Service**: This also cannot be verified from the document text, so always assign a 'warning'. The feedback must stress the importance of serving the schedule on the correct party at the address for notices stipulated in the contract, as improper service is equivalent to not serving one at all.
"""
    return ""

@app.post("/analyse-document")
async def analyse_document(doc_type: str = Form(...), file: UploadFile = File(...)):
    if not file:
        raise HTTPException(status_code=400, detail="No file uploaded.")

    try:
        # Read file content into a BytesIO object for processing
        file_content = await file.read()
        file_stream = io.BytesIO(file_content)
        
        extracted_text = extract_text_from_file(file_stream, file.filename)
        
        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="Could not extract any text from the document.")

        system_prompt = get_system_prompt(doc_type)
        if not system_prompt:
            raise HTTPException(status_code=400, detail="Invalid document type specified.")
            
        user_prompt = f"Here is the text from the document to be analysed:\n\n---DOCUMENT START---\n{extracted_text}\n---DOCUMENT END---"

        # Make the call to OpenAI API
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            response_format={"type": "json_object"}
        )
        
        ai_response_content = response.choices[0].message.content
        
        # The response should be a JSON string, which we parse and return
        return json.loads(ai_response_content)

    except HTTPException as e:
        # Re-raise HTTP exceptions to be handled by FastAPI
        raise e
    except Exception as e:
        print(f"ERROR in /analyse-document: {e}")
        return JSONResponse(content={"error": f"An unexpected error occurred: {str(e)}"}, status_code=500)


# -------------------------------------------------------------------
# --- START: LEXIFILE FUNCTIONALITY (MERGED & UPDATED) ---
# -------------------------------------------------------------------

def get_human_readable_type(content_type: str, filename: str) -> str:
    """Converts MIME type and filename to a user-friendly object type."""
    ext = filename.split('.')[-1].lower()
    if ext in ['msg', 'eml']:
        return "Email Message"
    if ext == 'pdf':
        return "PDF Document"
    if ext == 'docx':
        return "Word Document"
    if ext == 'doc':
        return "Legacy Word Document"
    if ext in ['png', 'jpg', 'jpeg', 'gif']:
        return "Image"
    return content_type or "Unknown File"

def get_renaming_system_prompt() -> str:
    """Returns the expert system prompt for the AI to rename documents."""
    return """
You are an expert legal discovery AI named LexiFile. Your task is to analyze text from a legal document and extract key information. Your entire response must be a single, valid JSON object.

The JSON object must have keys for:
1.  "date": The primary date (YYYY-MM-DD). Use today's date if none is found.
2.  "docType": The strict document type (e.g., "Email from [Sender's Email] to [Recipient's Email]", "Contract between [Party A] and [Party B]", "Affidavit of [Name]"). For emails, use the actual email addresses.
3.  "description": A brief 2-5 word summary of the subject matter.
4.  "keywords": An array of up to 10 relevant string keywords.
5.  "metadata": An object with the following keys. If info is not present, use an empty string "" or a sensible default.
    - "sender": The sender's email address ONLY if it's an email. Otherwise, this should be an empty string.
    - "recipient": The recipient's email address(es) ONLY if it's an email. Otherwise, this should be an empty string.
    - "author": The document author, extracted from document properties or text.
    - "lastModified": The last modified date (YYYY-MM-DD format).
"""

def extract_lexifile_text(file_path: str, filename: str) -> str:
    """Robustly extracts text from various file types from a file path."""
    text = ""
    lower_filename = filename.lower()
    try:
        if lower_filename.endswith('.pdf'):
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ""
            # OCR Fallback for scanned PDFs
            if len(text.strip()) < 50: 
                print(f"Performing OCR on {filename}...")
                text = pytesseract.image_to_string(file_path)
        elif lower_filename.endswith('.docx'):
            with open(file_path, "rb") as f:
                doc = docx.Document(f)
                core_props = doc.core_properties
                author = core_props.author if core_props.author else ""
                last_modified_by = core_props.last_modified_by if core_props.last_modified_by else ""
                text += f"[Author: {author}]\n[Last Modified By: {last_modified_by}]\n\n"
                for para in doc.paragraphs:
                    text += para.text + "\n"
        elif lower_filename.endswith('.doc'):
            try:
                text = pypandoc.convert_file(file_path, 'plain', format='doc')
            except Exception as pandoc_error:
                print(f"Pandoc failed for {filename}: {pandoc_error}. Falling back.")
                text = "" 
        elif lower_filename.endswith('.msg') or lower_filename.endswith('.eml'):
            msg = extract_msg.Message(file_path)
            text += f"From: {msg.sender}\nTo: {msg.to}\nCC: {msg.cc}\nSubject: {msg.subject}\n\n{msg.body}"
        return text
    except Exception as e:
        print(f"Error extracting text from {filename}: {e}")
        return ""

@app.post("/rename-document")
async def rename_document(file: UploadFile = File(...), project_id: str = Form(...)):
    global ARTIFACT_ID_COUNTER
    if not file or not project_id:
        raise HTTPException(status_code=400, detail="Missing file or project ID.")
    
    temp_path = os.path.join(LEXIFILE_STORAGE, file.filename)
    with open(temp_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    try:
        file_stat = os.stat(temp_path)
        last_modified_date = datetime.fromtimestamp(file_stat.st_mtime).strftime('%Y-%m-%d')
        
        extracted_text = extract_lexifile_text(temp_path, file.filename)
        
        ai_response = {}
        if not extracted_text.strip():
            doc_type = get_human_readable_type("", file.filename)
            ai_response = {
                "date": datetime.now().strftime('%Y-%m-%d'),
                "docType": doc_type, "description": "Media file",
                "keywords": ["media", file.filename.split('.')[-1]],
                "metadata": { "sender": "", "recipient": "", "author": "", "lastModified": last_modified_date }
            }
        else:
            system_prompt = get_renaming_system_prompt()
            user_prompt = f"Analyze the following text. Available metadata: Last Modified='{last_modified_date}'.\n\n---TEXT---\n{extracted_text[:8000]}\n---END---"
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
                response_format={"type": "json_object"}
            )
            ai_response = json.loads(response.choices[0].message.content)
            if not ai_response.get("metadata", {}).get("lastModified"):
                if "metadata" not in ai_response: ai_response["metadata"] = {}
                ai_response["metadata"]["lastModified"] = last_modified_date

        artifact_id = f"LEX-{str(ARTIFACT_ID_COUNTER).zfill(8)}"
        
        permanent_path = os.path.join(LEXIFILE_STORAGE, f"{artifact_id}_{file.filename}")
        os.rename(temp_path, permanent_path)

        date_part = ai_response.get('date', '0000-00-00').replace('-', '')
        doc_type_part = ai_response.get('docType', 'Untitled')
        description_part = ai_response.get('description', 'No Description')
        ext = file.filename.split('.')[-1]
        improved_filename = f"{date_part} - {doc_type_part} - {description_part}.{ext}"

        doc_record = {
            "artifactID": artifact_id, "projectID": project_id,
            "objectType": get_human_readable_type(file.content_type, file.filename),
            "originalFilename": file.filename,
            "improvedFilename": improved_filename,
            "date": ai_response.get('date', ''),
            "description": ai_response.get('description', ''),
            "author": ai_response.get('metadata', {}).get('author', ''),
            "lastModifiedOn": ai_response.get('metadata', {}).get('lastModified', ''),
            "sender": ai_response.get('metadata', {}).get('sender', ''),
            "recipient": ai_response.get('metadata', {}).get('recipient', ''),
            "uploadedBy": "demo.user@example.com",
            "keywords": ai_response.get('keywords', []),
            "ai_full_response": ai_response
        }
        
        DOCUMENTS_DB[artifact_id] = doc_record
        ARTIFACT_ID_COUNTER += 1
        
        return doc_record

    except Exception as e:
        print(f"Error in rename-document: {e}")
        if os.path.exists(temp_path):
            os.remove(temp_path)
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred: {e}")


@app.get("/get-receipt-url/{payment_intent_id}")
async def get_receipt_url(
    payment_intent_id: str,
    current_user: dict = Depends(get_current_purchase_user)
):
    """Gets the Stripe receipt URL for a payment intent"""
    try:
        # Verify this payment belongs to the user
        purchase = purchases_con.execute("""
            SELECT 1 FROM adjudicator_purchases 
            WHERE user_email = ? AND stripe_payment_intent_id = ?
        """, (current_user['email'], payment_intent_id)).fetchone()
        
        if not purchase:
            raise HTTPException(status_code=403, detail="Access denied")
        
        # Get the payment intent with expanded charge
        intent = stripe.PaymentIntent.retrieve(
            payment_intent_id,
            expand=['latest_charge']
        )
        
        receipt_url = None
        if intent.latest_charge:
            receipt_url = intent.latest_charge.receipt_url
        
        return {"receiptUrl": receipt_url}
        
    except stripe.error.StripeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        print(f"Get receipt URL error: {e}")
        raise HTTPException(status_code=500, detail="Failed to get receipt URL")

@app.get("/get-project-documents")
async def get_project_documents(project_id: str = Query(...)):
    if not project_id:
        raise HTTPException(status_code=400, detail="Project ID is required.")
    
    if project_id == 'all':
        return list(DOCUMENTS_DB.values())
        
    project_docs = [doc for doc in DOCUMENTS_DB.values() if doc.get("projectID") == project_id]
    return project_docs
    
@app.post("/fulltext-search")
async def fulltext_search(project_id: str = Form(...), query: str = Form(...)):
    if not query:
        return []

    results = []
    docs_to_search = [doc for doc in DOCUMENTS_DB.values() if project_id == 'all' or str(doc.get("projectID")) == project_id]
    
    for doc in docs_to_search:
        file_path = os.path.join(LEXIFILE_STORAGE, f"{doc['artifactID']}_{doc['originalFilename']}")
        if os.path.exists(file_path):
            content = extract_lexifile_text(file_path, doc['originalFilename'])
            if query.lower() in content.lower():
                match_pos = content.lower().find(query.lower())
                start = max(0, match_pos - 25)
                end = min(len(content), match_pos + len(query) + 25)
                snippet = "..." + content[start:end].strip().replace("\n", " ") + "..."
                doc_with_snippet = doc.copy()
                doc_with_snippet["snippet"] = snippet
                results.append(doc_with_snippet)
    return results

@app.get("/download-single-doc/{artifact_id}")
async def download_single_doc(artifact_id: str):
    doc_record = DOCUMENTS_DB.get(artifact_id)
    if not doc_record:
        raise HTTPException(status_code=404, detail="Document not found")
    
    file_path = os.path.join(LEXIFILE_STORAGE, f"{doc_record['artifactID']}_{doc_record['originalFilename']}")
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found on server")

    return FileResponse(file_path, filename=doc_record['improvedFilename'])

@app.post("/bulk-download")
async def bulk_download(artifact_ids: List[str] = Body(...)):
    def files_generator():
        z = zipstream.ZipFile(mode='w', compression=zipstream.ZIP_DEFLATED)
        for artifact_id in artifact_ids:
            doc_record = DOCUMENTS_DB.get(artifact_id)
            if doc_record:
                file_path = os.path.join(LEXIFILE_STORAGE, f"{doc_record['artifactID']}_{doc_record['originalFilename']}")
                if os.path.exists(file_path):
                    z.write(file_path, arcname=doc_record['improvedFilename'])
        for chunk in z:
            yield chunk

    return StreamingResponse(files_generator(), media_type="application/zip", headers={
        'Content-Disposition': f'attachment; filename="LexiFile_Bulk_Export_{datetime.now().strftime("%Y%m%d")}.zip"'
    })
    
@app.delete("/delete-document/{artifact_id}")
async def delete_document(artifact_id: str):
    if artifact_id in DOCUMENTS_DB:
        doc_record = DOCUMENTS_DB[artifact_id]
        file_path = os.path.join(LEXIFILE_STORAGE, f"{doc_record['artifactID']}_{doc_record['originalFilename']}")
        if os.path.exists(file_path):
            os.remove(file_path)
        del DOCUMENTS_DB[artifact_id]
        return {"status": "success", "message": "Document deleted"}
    raise HTTPException(status_code=404, detail="Document not found")


@app.post("/preview-email")
async def preview_email(file: UploadFile = File(...)):
    if not file or not (file.filename.lower().endswith('.msg') or file.filename.lower().endswith('.eml')):
        raise HTTPException(status_code=400, detail="Invalid file type for email preview.")
    try:
        file_content = await file.read()
        msg = extract_msg.Message(io.BytesIO(file_content))
        return {
            "from": msg.sender, "to": msg.to, "cc": msg.cc,
            "subject": msg.subject, "date": msg.date,
            "body": msg.body
        }
    except Exception as e:
        print(f"Error parsing email file {file.filename}: {e}")
        raise HTTPException(status_code=500, detail="Failed to parse email file.")

# --- Project Management Endpoints ---
@app.get("/projects")
async def get_projects():
    # Add document count to each project
    for project in PROJECTS_DB:
        count = sum(1 for doc in DOCUMENTS_DB.values() if str(doc.get("projectID")) == str(project["id"]))
        project["documentCount"] = count
    return PROJECTS_DB

@app.post("/create-project")
async def create_project(projectName: str = Form(...), clientName: str = Form(...), matterNumber: str = Form(...)):
    project_id = len(PROJECTS_DB) + 1
    new_project = {
        "id": project_id, "name": projectName, "client": clientName,
        "matter": matterNumber, "dateCreated": datetime.now().strftime("%Y-%m-%d")
    }
    PROJECTS_DB.append(new_project)
    return new_project

# -------------------------------------------------------------------
# --- END: LEXIFILE FUNCTIONALITY ---
# -------------------------------------------------------------------


# ---------- DB Download ----------
@app.get("/download-db")
async def download_db():
    return FileResponse("/tmp/qbcc.db", filename="qbcc.db")









# ---------------- STRIPE & PURCHASE LOGIC ----------------
stripe.api_key = os.getenv("STRIPE_SECRET_KEY")

# In server.py, REPLACE your existing /check-adjudicator-access function with this one

@app.get("/check-adjudicator-access/{adjudicator_name}")
def check_adjudicator_access(
    adjudicator_name: str,
    current_user: dict = Depends(get_current_purchase_user)
):
    """
    Checks if the user has access, either through full privileges or a specific purchase.
    """
    try:
        # --- NEW: First, check for full access privileges ---
        full_access = purchases_con.execute("""
            SELECT 1 FROM full_access_users WHERE email = ?
        """, (current_user["email"],)).fetchone()

        if full_access:
            return {"hasAccess": True}
        # --- END NEW ---

        decoded_name = unquote_plus(adjudicator_name)
        
        # Original logic: check for a specific purchase record
        purchase = purchases_con.execute("""
            SELECT 1 FROM adjudicator_purchases 
            WHERE user_email = ? AND adjudicator_name = ?
        """, (current_user["email"], decoded_name)).fetchone()
        
        has_access = purchase is not None
        return {"hasAccess": has_access}
        
    except Exception as e:
        print(f"Error checking access for {current_user['email']} to {adjudicator_name}: {e}")
        return {"hasAccess": False}
    
# In server (43).py

@app.post("/create-payment-intent")
async def create_payment_intent(
    adjudicator_name: str = Form(...),
    current_user: dict = Depends(get_current_purchase_user)
):
    """Creates a simple Stripe PaymentIntent. Invoice created after successful payment."""
    try:
        if not stripe.api_key:
            raise HTTPException(status_code=500, detail="Stripe API key is not configured.")

        # Step 1: Find or Create a Stripe Customer
        customers = stripe.Customer.list(email=current_user['email'], limit=1)
        if customers.data:
            customer = customers.data[0]
        else:
            customer = stripe.Customer.create(
                email=current_user['email'],
                name=f"{current_user.get('first_name', '')} {current_user.get('last_name', '')}"
            )

        # Step 2: Create a simple PaymentIntent (no invoice yet)
        intent = stripe.PaymentIntent.create(
            amount=6995,
            currency='aud',
            customer=customer.id,
            description=f"Access to Adjudicator Insights: {adjudicator_name}",
            metadata={
                'user_email': current_user['email'],
                'adjudicator_name': adjudicator_name
            }
        )

        return {
            'clientSecret': intent.client_secret,
            'paymentIntentId': intent.id
        }
    except stripe.error.StripeError as e:
        print(f"Stripe error: {e}")
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        print(f"Stripe PaymentIntent creation failed: {e}")
        raise HTTPException(status_code=400, detail=str(e))
    
@app.post("/confirm-purchase")
async def confirm_purchase(
    payment_intent_id: str = Form(...),
    adjudicator_name: str = Form(...),
    current_user: dict = Depends(get_current_purchase_user)
):
    """Confirms purchase after successful payment."""
    try:
        # Verify the payment succeeded - EXPAND charges to get receipt URL
        intent = stripe.PaymentIntent.retrieve(
            payment_intent_id,
            expand=['latest_charge']  # This expands the charge data
        )
        
        if intent.status != 'succeeded':
            raise HTTPException(status_code=400, detail="Payment not completed")

        # Check if purchase already exists
        existing = purchases_con.execute("""
            SELECT 1 FROM adjudicator_purchases 
            WHERE stripe_payment_intent_id = ?
        """, (payment_intent_id,)).fetchone()
        
        if existing:
            return {
                "msg": "Purchase already recorded",
                "adjudicator_name": adjudicator_name,
                "receiptUrl": None
            }
        
        # Record the purchase
        purchases_con.execute("""
            INSERT INTO adjudicator_purchases
            (user_email, adjudicator_name, stripe_payment_intent_id, amount_paid)
            VALUES (?, ?, ?, ?)
        """, (current_user['email'], adjudicator_name, payment_intent_id, 69.95))
        purchases_con.commit()
        
        # Get Stripe's built-in receipt URL from the expanded charge
        receipt_url = None
        if intent.latest_charge:
            receipt_url = intent.latest_charge.receipt_url
        
        return {
            "msg": "Purchase confirmed",
            "adjudicator_name": adjudicator_name,
            "receiptUrl": receipt_url
        }

    except stripe.error.StripeError as e:
        print(f"Stripe error in confirm-purchase: {e}")
        raise HTTPException(status_code=400, detail=f"Stripe error: {str(e)}")
    except Exception as e:
        print(f"Purchase confirmation error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Failed to confirm purchase")
    

@app.post("/purchase-register")
def purchase_register(
    email: str = Form(...),
    password: str = Form(...),
    first_name: str = Form(...),
    last_name: str = Form(...),
    firm_name: str = Form(""), # Add default value
    abn: str = Form(""),       # Add default value
    billing_address: str = Form(...),
    billing_city: str = Form(...),
    billing_state: str = Form(...),
    billing_postcode: str = Form(...),
    employee_size: str = Form(""), # Add default value
    phone: str = Form("")
):

    """Enhanced registration with company details"""
    try:
        hashed_password = get_password_hash(password)
        
        purchases_con.execute("""
            INSERT INTO purchase_users (
                email, hashed_password, first_name, last_name, 
                firm_name, abn, billing_address, billing_city, 
                billing_state, billing_postcode, employee_size, phone
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            email, hashed_password, first_name, last_name,
            firm_name, abn, billing_address, billing_city,
            billing_state, billing_postcode, employee_size, phone
        ))
        purchases_con.commit()
        
        return {"msg": "Registration successful"}
        
    except sqlite3.IntegrityError:
        raise HTTPException(status_code=400, detail="Email already registered")
    except Exception as e:
        print(f"Registration error: {e}")
        raise HTTPException(status_code=500, detail="Registration failed")


@app.post("/purchase-login")
def purchase_login(form_data: OAuth2PasswordRequestForm = Depends()):
    """Enhanced login with last_login tracking"""
    user = get_purchase_user_by_email(form_data.username)
    if not user or not verify_password(form_data.password, user["hashed_password"]):
        raise HTTPException(status_code=401, detail="Incorrect email or password")
    
    # Update last login time
    purchases_con.execute(
        "UPDATE purchase_users SET last_login = CURRENT_TIMESTAMP WHERE email = ?",
        (user["email"],)
    )
    purchases_con.commit()
    
    access_token = create_access_token(
        data={"sub": user["email"]},
        expires_delta=timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    )
    return {"access_token": access_token, "token_type": "bearer"}
# ... surrounding code ...


# In server.py, REPLACE your existing /purchase-me function with this one

@app.get("/purchase-me")
def read_purchase_user_me(current_user: dict = Depends(get_current_purchase_user)):
    """Get current purchase user info including full access status."""

    # --- NEW: Check for full access privileges ---
    full_access_check = purchases_con.execute(
        "SELECT 1 FROM full_access_users WHERE email = ?",
        (current_user["email"],)
    ).fetchone()
    has_full_access = full_access_check is not None
    # --- END NEW ---

    user_data = {
        "email": current_user["email"],
        "first_name": current_user["first_name"],
        "last_name": current_user["last_name"],
        "firm_name": current_user["firm_name"],
        "abn": current_user["abn"],
        "billing_address": current_user["billing_address"],
        "billing_city": current_user["billing_city"],
        "billing_state": current_user["billing_state"],
        "billing_postcode": current_user["billing_postcode"],
        "employee_size": current_user["employee_size"],
        "phone": current_user["phone"],
        "created_at": current_user["created_at"],
        "last_login": current_user.get("last_login"),
        "email_verified": current_user.get("email_verified", 0),
        "has_full_access": has_full_access # ADDED THIS LINE
    }
    return user_data



@app.put("/update-profile")
def update_profile(
    current_user: dict = Depends(get_current_purchase_user),
    first_name: str = Form(...),
    last_name: str = Form(...),
    firm_name: str = Form(...),
    abn: str = Form(...),
    billing_address: str = Form(...),
    billing_city: str = Form(...),
    billing_state: str = Form(...),
    billing_postcode: str = Form(...),
    employee_size: str = Form(...),
    phone: str = Form("")
):
    """Update user profile information"""
    try:
        purchases_con.execute("""
            UPDATE purchase_users SET
                first_name = ?, last_name = ?, firm_name = ?,
                abn = ?, billing_address = ?, billing_city = ?,
                billing_state = ?, billing_postcode = ?, employee_size = ?, phone = ?
            WHERE email = ?
        """, (
            first_name, last_name, firm_name, abn, billing_address,
            billing_city, billing_state, billing_postcode, employee_size,
            phone, current_user["email"]
        ))
        purchases_con.commit()
        return {"msg": "Profile updated successfully"}
    except Exception as e:
        print(f"Profile update error: {e}")
        raise HTTPException(status_code=500, detail="Failed to update profile")


@app.get("/my-purchases")
def get_my_purchases(current_user: dict = Depends(get_current_purchase_user)):
    """Fetches all purchase records for the currently authenticated user."""
    try:
        purchases = purchases_con.execute("""
            SELECT adjudicator_name, purchase_date, amount_paid, 
                   stripe_payment_intent_id, stripe_invoice_id
            FROM adjudicator_purchases
            WHERE user_email = ?
            ORDER BY purchase_date DESC
        """, (current_user["email"],)).fetchall()
        
        result = [dict(row) for row in purchases]
        print(f"Returning purchases: {result}")  # Debug log
        return result
        
    except Exception as e:
        print(f"Error fetching purchases for {current_user['email']}: {e}")
        raise HTTPException(status_code=500, detail="Could not retrieve purchase history.")
    

@app.get("/download-invoice/{invoice_id}")
async def download_invoice(
    invoice_id: str,
    current_user: dict = Depends(get_current_purchase_user)
):
    """Securely streams a Stripe invoice PDF to an authenticated user."""
    try:
        # Retrieve the invoice from Stripe
        invoice = stripe.Invoice.retrieve(invoice_id)

        # Security Check: Ensure the invoice belongs to the current user
        customers = stripe.Customer.list(email=current_user['email'], limit=1)
        if not customers.data or customers.data[0].id != invoice.customer:
            raise HTTPException(status_code=403, detail="Permission denied")

        # Get the PDF URL from the invoice object
        invoice_pdf_url = invoice.invoice_pdf
        if not invoice_pdf_url:
            raise HTTPException(status_code=404, detail="Invoice PDF not available yet.")

        # Stream the PDF content
        response = requests.get(invoice_pdf_url)
        response.raise_for_status() # Raise an exception for bad status codes

        return StreamingResponse(
            io.BytesIO(response.content),
            media_type='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="Sopal_Invoice_{invoice.number}.pdf"'
            }
        )
    except stripe.error.StripeError as e:
        raise HTTPException(status_code=404, detail=f"Stripe error: {e.user_message}")
    except Exception as e:
        print(f"Invoice download error: {e}")
        raise HTTPException(status_code=500, detail="Failed to download invoice.")
    
@app.get("/payment-methods")
async def get_payment_methods(current_user: dict = Depends(get_current_purchase_user)):
    """Get user's saved payment methods from Stripe"""
    try:
        # Get or create Stripe customer
        customers = stripe.Customer.list(email=current_user['email'], limit=1)
        
        if customers.data:
            customer = customers.data[0]
        else:
            # Create new customer if doesn't exist
            customer = stripe.Customer.create(
                email=current_user['email'],
                name=f"{current_user['first_name']} {current_user['last_name']}"
            )
        
        # Get payment methods
        payment_methods = stripe.PaymentMethod.list(
            customer=customer.id,
            type='card'
        )
        
        return {
            "customer_id": customer.id,
            "payment_methods": [
                {
                    "id": pm.id,
                    "brand": pm.card.brand,
                    "last4": pm.card.last4,
                    "exp_month": pm.card.exp_month,
                    "exp_year": pm.card.exp_year
                } for pm in payment_methods.data
            ]
        }
        
    except Exception as e:
        print(f"Error fetching payment methods: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch payment methods")


@app.post("/setup-intent")
async def create_setup_intent(current_user: dict = Depends(get_current_purchase_user)):
    """Creates a Stripe SetupIntent for adding payment methods"""
    try:
        customers = stripe.Customer.list(email=current_user['email'], limit=1)
        
        if customers.data:
            customer = customers.data[0]
        else:
            customer = stripe.Customer.create(
                email=current_user['email'],
                name=f"{current_user['first_name']} {current_user['last_name']}"
            )
        
        intent = stripe.SetupIntent.create(
            customer=customer.id,
            payment_method_types=['card']
        )
        
        return {"client_secret": intent.client_secret}
        
    except Exception as e:
        print(f"Error creating setup intent: {e}")
        raise HTTPException(status_code=400, detail=str(e))


@app.delete("/payment-method/{payment_method_id}")
async def delete_payment_method(
    payment_method_id: str,
    current_user: dict = Depends(get_current_purchase_user)
):
    """Deletes a saved payment method"""
    try:
        stripe.PaymentMethod.detach(payment_method_id)
        return {"msg": "Payment method removed"}
    except Exception as e:
        print(f"Error deleting payment method: {e}")
        raise HTTPException(status_code=400, detail=str(e))



@app.get("/api/decision-text/{decision_id}")
def get_decision_text(decision_id: str = Path(...)):
    """Get full text for a specific decision"""
    try:
        row = con.execute(
            "SELECT full_text FROM docs_fresh WHERE ejs_id = ?", 
            (decision_id,)
        ).fetchone()
        
        if not row:
            raise HTTPException(status_code=404, detail="Decision not found")
        
        return {"fullText": row["full_text"] or ""}
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error fetching decision text: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch decision text")

    


from io import BytesIO
from datetime import datetime
from fastapi import Response, Header, HTTPException
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors


from io import BytesIO
from datetime import datetime
from fastapi import Response, Header, HTTPException
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors



# --- Helper for summary data ---
def _get_adjudicator_summary_data(db_con, adjudicator_name: str) -> dict:
    """Internal helper to fetch summary data for a single adjudicator."""
    query = """
        SELECT 
            name,
            totalDecisions,
            totalClaimAmount,
            totalAwardedAmount,
            avgAwardRate,
            avgClaimantFeeProportion,
            avgRespondentFeeProportion,
            zeroAwardCount
        FROM (
            SELECT 
                a.adjudicator_name as name,
                COUNT(*) as totalDecisions,
                SUM(CASE WHEN a.claimed_amount NOT IN ('N/A', '') AND a.claimed_amount IS NOT NULL THEN CAST(a.claimed_amount AS REAL) ELSE 0 END) as totalClaimAmount,
                SUM(CASE WHEN a.adjudicated_amount NOT IN ('N/A', '') AND a.adjudicated_amount IS NOT NULL THEN CAST(a.adjudicated_amount AS REAL) ELSE 0 END) as totalAwardedAmount,
                AVG(CASE WHEN CAST(a.claimed_amount AS REAL) > 0 THEN (CAST(a.adjudicated_amount AS REAL) * 100.0 / CAST(a.claimed_amount AS REAL)) ELSE 0 END) as avgAwardRate,
                AVG(CAST(a.fee_claimant_proportion AS REAL)) as avgClaimantFeeProportion,
                AVG(CAST(a.fee_respondent_proportion AS REAL)) as avgRespondentFeeProportion,
                SUM(CASE WHEN a.adjudicated_amount = '0' OR a.adjudicated_amount = '0.0' THEN 1 ELSE 0 END) as zeroAwardCount
            FROM decision_details a
            WHERE LOWER(TRIM(a.adjudicator_name)) = LOWER(TRIM(?))
            GROUP BY a.adjudicator_name
        )
    """
    row = db_con.execute(query, (adjudicator_name,)).fetchone()
    return dict(row) if row else {}

@app.post("/generate-summary-pdf")
async def generate_summary_pdf(
    adjudicator_name: str = Form(...),
    current_user: dict = Depends(get_current_purchase_user)
):
    """
    Generates a branded PDF summary of an adjudicator's insights for logged-in users.
    """
    # Step 1: Get adjudicator data using our reliable helper
    data = _get_adjudicator_summary_data(con, adjudicator_name)
    if not data:
        raise HTTPException(status_code=404, detail="Adjudicator not found")

    decisions_rows = con.execute("""
        SELECT claimant_name, respondent_name, decision_date, claimed_amount, adjudicated_amount, 
               fee_claimant_proportion, fee_respondent_proportion, pdf_path
        FROM decision_details
        WHERE LOWER(TRIM(adjudicator_name)) = LOWER(TRIM(?))
        ORDER BY decision_date DESC
    """, (adjudicator_name,)).fetchall()

    decisions = [dict(row) for row in decisions_rows]
    data["decisions"] = decisions

    # Step 2: Build the PDF
    buffer = BytesIO()
    pdf = SimpleDocTemplate(
        buffer, pagesize=A4, topMargin=40, bottomMargin=40, leftMargin=45, rightMargin=45
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Heading", fontName="Helvetica-Bold", fontSize=14, textColor=colors.HexColor("#00C97C")))
    styles.add(ParagraphStyle(name="Label", fontName="Helvetica-Bold", fontSize=10, textColor=colors.black))
    styles.add(ParagraphStyle(name="Body", fontName="Helvetica", fontSize=10, leading=14))
    styles.add(ParagraphStyle(name="Footer", fontName="Helvetica-Oblique", fontSize=8, textColor=colors.gray, alignment=1))

    story = []

    # --- Header ---
    story.append(Paragraph("<b>Search adjudicator's decisions</b>", styles["Heading"]))
    story.append(Paragraph("www.sopal.com.au &nbsp;&nbsp;&nbsp;&nbsp; info@sopal.com.au", styles["Body"]))
    story.append(Spacer(1, 12))

    export_no = f"SJE{datetime.now():%Y%m%d}{str(current_user['id']).zfill(6)}"
    story.append(Paragraph(f"<b>Export No:</b> {export_no}", styles["Body"]))
    story.append(Paragraph(f"<b>Export for:</b> {current_user.get('first_name')} {current_user.get('last_name')}", styles["Body"]))
    if current_user.get("firm_name"):
        story.append(Paragraph(current_user.get("firm_name"), styles["Body"]))
    story.append(Paragraph(
        f"{current_user.get('billing_address', '')}, {current_user.get('billing_city', '')} "
        f"{current_user.get('billing_state', '')} {current_user.get('billing_postcode', '')}",
        styles["Body"]
    ))
    story.append(Spacer(1, 18))

    # --- Adjudicator Summary Table ---
    story.append(Paragraph("<b>Adjudicator Insights Summary</b>", styles["Heading"]))
    story.append(Spacer(1, 6))

    total_decisions = data.get("totalDecisions", 1)
    summary_data = [
        ["Adjudicator’s Name", data.get("name", "")],
        ["Number of Decisions", str(data.get("totalDecisions", 0))],
        ["Total Claimed Amount", f"${data.get('totalClaimAmount', 0):,.0f}"],
        ["Total Adjudicated Amount", f"${data.get('totalAwardedAmount', 0):,.0f}"],
        ["Avg Claimant Fee Proportion", f"{data.get('avgClaimantFeeProportion', 0):.1f}%"],
        ["Avg Respondent Fee Proportion", f"{data.get('avgRespondentFeeProportion', 0):.1f}%"],
        ["Number of Nil Decisions", str(data.get('zeroAwardCount', 0))],
        ["Avg Portion Awarded Per Decision", f"{data.get('avgAwardRate', 0):.1f}%"],
    ]

    table = Table(summary_data, colWidths=[200, 200])
    # ... (rest of the table styling, which is correct) ...
    table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#F8F9FA")),
        ('FONT', (0, 0), (-1, -1), 'Helvetica', 9),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(table)
    story.append(Spacer(1, 18))

    # --- Decision History ---
    story.append(Paragraph("<b>Adjudicator Decision History</b>", styles["Heading"]))
    story.append(Spacer(1, 8))

    for d in data.get("decisions", []):
        story.append(Paragraph(f"<b>{d.get('claimant_name', '')} v {d.get('respondent_name', '')}</b>", styles["Label"]))
        story.append(Paragraph(f"Date: {d.get('decision_date', 'N/A')} | Claimed: ${float(d.get('claimed_amount', 0) or 0):,.0f} | Adjudicated: ${float(d.get('adjudicated_amount', 0) or 0):,.0f}", styles["Body"]))
        story.append(Spacer(1, 8))

    pdf.build(story)
    buffer.seek(0)

    headers = {"Content-Disposition": f"attachment; filename={adjudicator_name}_SopalInsights.pdf"}
    return Response(buffer.read(), media_type="application/pdf", headers=headers)


@app.post("/create-invoice-for-purchase")
async def create_invoice_for_purchase(
    adjudicator_name: str = Form(...),
    payment_intent_id: str = Form(...),
    current_user: dict = Depends(get_current_purchase_user)
):
    """Creates a Stripe Invoice for an existing purchase."""
    try:
        print(f"Looking for purchase: user={current_user['email']}, payment_intent={payment_intent_id}")
        
        # Get the purchase record
        purchase = purchases_con.execute("""
            SELECT * FROM adjudicator_purchases 
            WHERE user_email = ? AND stripe_payment_intent_id = ?
        """, (current_user['email'], payment_intent_id)).fetchone()
        
        print(f"Purchase found: {purchase}")
        
        if not purchase:
            # Let's see what purchases exist for this user
            all_purchases = purchases_con.execute("""
                SELECT stripe_payment_intent_id FROM adjudicator_purchases 
                WHERE user_email = ?
            """, (current_user['email'],)).fetchall()
            print(f"All payment_intent_ids for user: {[dict(p) for p in all_purchases]}")
            raise HTTPException(status_code=404, detail="Purchase not found")
        
        # Check if invoice already exists
        if purchase['stripe_invoice_id']:
            return {"invoiceId": purchase['stripe_invoice_id']}
        
        # Get or create Stripe customer
        customers = stripe.Customer.list(email=current_user['email'], limit=1)
        if customers.data:
            customer = customers.data[0]
        else:
            customer = stripe.Customer.create(
                email=current_user['email'],
                name=f"{current_user.get('first_name', '')} {current_user.get('last_name', '')}"
            )
        
        # Create invoice item
        stripe.InvoiceItem.create(
            customer=customer.id,
            amount=6995,
            currency='aud',
            description=f"Adjudicator Insights: {adjudicator_name}"
        )
        
        # Create and finalize invoice
        invoice = stripe.Invoice.create(
            customer=customer.id,
            auto_advance=True,
            collection_method='send_invoice',
            days_until_due=0,
            metadata={
                'user_email': current_user['email'],
                'adjudicator_name': adjudicator_name,
                'payment_intent_id': payment_intent_id
            }
        )
        
        invoice = stripe.Invoice.finalize_invoice(invoice.id)
        invoice = stripe.Invoice.pay(invoice.id, paid_out_of_band=True)
        
        # Store invoice ID
        purchases_con.execute("""
            UPDATE adjudicator_purchases 
            SET stripe_invoice_id = ? 
            WHERE stripe_payment_intent_id = ?
        """, (invoice.id, payment_intent_id))
        purchases_con.commit()
        
        return {"invoiceId": invoice.id}
        
    except stripe.error.StripeError as e:
        print(f"Stripe error creating invoice: {e}")
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        print(f"Error creating invoice: {e}")
        raise HTTPException(status_code=500, detail="Failed to create invoice")

@app.post("/api/adjudicator-decisions-text")
async def get_adjudicator_decisions_text(decision_ids: List[str] = Body(...)):
    """Batch fetch full text for multiple decisions"""
    try:
        if not decision_ids:
            return {"decisions": []}
        
        # Create placeholders for SQL IN clause
        placeholders = ','.join('?' * len(decision_ids))
        
        rows = con.execute(
            f"SELECT ejs_id, full_text FROM docs_fresh WHERE ejs_id IN ({placeholders})",
            decision_ids
        ).fetchall()
        
        result = {row["ejs_id"]: row["full_text"] or "" for row in rows}
        
        return {"decisions": result}
        
    except Exception as e:
        print(f"Error batch fetching decision texts: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch decision texts")


# ==================== BIF ACT AI CHATBOT ====================

# Knowledge base content - the three requirements from your knowledge centre
BIF_ACT_KNOWLEDGE_BASE = """
# BIF Act Payment Claim Requirements

## 1. IDENTIFY THE CONSTRUCTION WORKS

Pursuant to section 68(1)(a) of the BIF Act, a payment claim must identify the construction work or related goods and services to which it relates.

### General Principles
The test is not whether the claim explains every calculation, but whether the relevant construction work (and/or related goods and services) is sufficiently identified such that the basis of the claim is reasonably comprehensible to the recipient.

Key factors:
- Payment claims are given and received by people experienced in the building industry
- The parties' background knowledge from past dealings is relevant
- The claim may reference the construction contract to identify items
- A practical, not overly technical attitude is taken

### Degree of Particularity
The construction work must be identified with some degree of particularity. The specific work should be identified, not just the general work of the whole contract. The respondent should be capable of determining whether to pay the claim.

However, a few failures to identify work won't necessarily invalidate the claim if most work is properly identified.

### What Fails
- Percentages against broad trade headings without tethering to identifiable scopes/locations/activities (e.g., "5% of concreting" for a 56-townhouse project)
- Generic categories with only percentages and no particulars
- Trade breakdowns like "Electrical", "Mechanical", "Builder's Preliminaries" without specifics

### Practice Tips
- Use location + scope + activity + period for every line (e.g., "Tower 1, L10–L12: install grid & tile to corridors – 1 January 2025 to 10 January 2025")
- Anchor to the contract: cite contract/Annexure item ID, unit, quantity, rate, previously approved and this period columns
- Avoid percentages against broad trades; make it findable on site with concrete descriptors
- Keep all supporting documents with the claim
- For supply-only items, specify units/tonnage/quantities, delivery dates/locations
- Variations should ideally be communicated by project letter first, then attached to the claim

## 2. STATE THE CLAIMED AMOUNT

Pursuant to section 68(1)(b) of the BIF Act, a payment claim must state the amount of the progress payment (i.e. the claimed amount) that the claimant claims is payable by the respondent.

### Must State a Claimed Amount
To avoid ambiguity, a payment claim should clearly identify a single, specific figure as the "claimed amount". Failing to state the claimed amount may invalidate the payment claim, even if the total is ascertainable by summing individual items.

### Avoid Multiple Figures
A payment claim should avoid presenting multiple figures that could each be read as the "claimed amount". Multiple spreadsheets with different "due this claim" totals for separable parts, with no single total, will fail.

However, there is authority suggesting multiple amounts which collectively constitute the claimed amount may be sufficient, as long as they're stated as part of one payment claim.

### GST Inclusive
To the extent GST applies, best practice is to express the claimed amount as the GST inclusive figure. Expressing it as GST exclusive may expose the claim to argument that the GST component is not part of the payment claim.

## 3. REQUEST PAYMENT

Pursuant to section 68(1)(c) of the BIF Act, a payment claim must request payment of the claimed amount.

### "Invoice" is Sufficient
It is sufficient to include the word "invoice" on the face of the payment claim (section 68(3) of the BIF Act). However, best practice is to ensure it's a valid tax invoice.

### What's Not Sufficient
The words "amount due this claim" is NOT sufficient. Something more is required - there must be a request for payment of the amount claimed.

### Must Be Express or Clearly Implied
A request for payment must be either expressly stated in the payment claim, or necessarily and clearly implied. 

Simply stating the payment claim is made under the BIF Act is not, by itself, sufficient. However, a document with "payment claim" and stating "This is a payment claim made under the Building Industry Fairness (Security of Payment) Act 2017" may be valid, especially with additional language like "net payment claimed".

### Practice Tips
- Include the word "invoice" or a phrase like "this is a request for payment under the Building Industry Fairness (Security of Payment) Act 2017 (Qld)" (or both)
- If relying only on "invoice", ensure the document complies with tax invoice requirements (claimant's name and ABN, taxable supplies identified, GST amounts, date of issue)

## 4. CLAIM FOR A VALID REFERENCE DATE

Pursuant to section 70 of the BIF Act, a person who has carried out construction work, or supplied related goods and services, under a construction contract is entitled to a progress payment from each reference date.

A "reference date" is defined under section 67 of the BIF Act as:
- (a) the date stated in, or worked out under, the contract as the date on which a progress payment may be made; or
- (b) if the contract does not provide for the matter, the last day of the month.

### Jurisdictional Requirement
A payment claim will be invalid if it is not made from a valid reference date. As the High Court unanimously found in Southern Han Breakfast Point Pty Ltd (in liq) v Lewence Construction Pty Ltd (2016) 260 CLR 340:

The existence of a reference date under a construction contract is a precondition to the making of a valid payment claim. The High Court confirmed that where one payment claim has already been made in relation to a particular reference date, any later payment claim purportedly made in respect of the same reference date is not valid.

Thus, if a payment claim is made without the accrual of a valid reference date, then it is invalid and cannot be adjudicated upon.

### From Each Reference Date
Under the BIF Act, the correct approach is that a payment claim can be served both on and from/after the reference date. A claim can be made at any time after the contractual date, subject to the statutory outer limit (6 months).

### Cannot Claim Twice for the Same Reference Date
Pursuant to section 75(4) of the BIF Act, a claimant cannot make more than one payment claim for each reference date. If a payment claim is made in respect of a reference date already relied upon for a previous valid payment claim, then it will be invalid.

### Work Must Be Up To Reference Date
There is authority in Queensland which suggests that a payment claim will be invalid if it claims for work performed after the reference date:
- Karam Group Pty Ltd v Earthmoving Contractors Pty Ltd [2021] QSC 10 stated that a payment claim which includes work beyond the reference date is not a valid payment claim
- Watkins Contracting Pty Ltd v Hyatt Ground Engineering Pty Ltd [2018] QSC 065 held that payment claims which included work carried out after the reference date were not valid

The reference date is the point in time at which a payment claim is to be assessed. Events which post-date the reference date are not strictly relevant to the valuation of the payment claim.

### Contractual Preconditions
Generally, a contractual provision that purports to set preconditions to the accrual of a reference date will be void for the purposes of the BIF Act. A provision which inordinately delays or effectively prevents a reference date from arising would be contrary to the Act's objective.

### Reference Dates Following Take Out
Under many unamended Australian Standard contracts, payment to the claimant is suspended upon the respondent validly taking the work out of its hands. In such cases, reference dates do not arise upon payments being suspended following take out.

### Reference Dates Following Termination
If a construction contract is terminated and the contract does not provide for, or purports to prevent, a reference date surviving beyond termination, the final reference date for the contract is the date the contract is terminated.

### Practice Tips
- Serve only on or after the reference date, never before
- Make one claim per reference date
- If needed, re-include unpaid amounts in a later claim rather than issuing a second claim for the same date
- Include only work up to the reference date (exclude work done later, late variations, or post-date quantities)
- On termination/take-out, confirm whether reference dates survive

## 5. NOT BE FOR UNLICENSED BUILDING WORK

Under section 42 of the Queensland Building and Construction Commission Act 1991 (Qld), a person who performs unlicensed building work is not entitled to any monetary or other consideration – including under the BIF Act – for doing so.

Section 42(3) states: "A person who carries out building work in contravention of this section is not entitled to any monetary or other consideration for doing so."

### Jurisdictional Bar
A payment claim cannot be made for construction work carried out by an unlicensed person if the QBCC Act requires that person to be licensed to carry out that work. This is a common jurisdictional challenge in adjudication.

### Contract Cannot Be Severed
In Dart Holdings Pty Ltd v Total Concept Group Pty Ltd [2012] QSC 158, the Court held that where part of a lump-sum contract required unlicensed work (glass awnings outside the carpentry licence), the entire contract was tainted. Because the contract could not be severed to save part of the payment provisions, there was no security of payment entitlement and the adjudicator's decision was void.

### Small Items Can Taint Large Contracts
In Galaxy Developments Pty Ltd v Civil Contractors (Aust) Pty Ltd [2020] QSC 51, a $1.3 million contract was held unenforceable because the claimant lacked appropriate licences for minor "bus stop" items (a freestanding seat and bike rack). The Court acknowledged this was "absurd in reality" but correct in law, highlighting the harsh consequences of licensing requirements.

### Exemptions Apply
In Ingeteam Australia Pty Ltd v Susan River Solar Pty Ltd [2024] QSC 30, a $2.39m payment claim was initially rejected because of a $294 plywood "floor repair". However, the Court held the adjudicator had misapplied the "floating floor" and sub-$3,300 exemption, so no licence was actually required.

### Practice Tips
- As a claimant, do a licence-scope audit for every item/variation
- If using the head-contractor exemption, require and file proof that only appropriately licensed subcontractors will perform building work
- Respondents should perform QBCC licence searches to check whether the claimant held appropriate licence(s) at all material times
"""

BIF_ACT_SYSTEM_PROMPT = f"""You are an expert AI assistant specialized in Queensland's Building Industry Fairness (Security of Payment) Act 2017 (BIF Act), specifically regarding payment claim requirements.

Your knowledge base covers five critical requirements for valid payment claims:
1. Identify the construction works
2. State the claimed amount
3. Request payment
4. Claim for a valid reference date
5. Not be for unlicensed building work

{BIF_ACT_KNOWLEDGE_BASE}

When answering questions:
- Be precise and cite specific requirements from the BIF Act
- Provide practical, actionable advice
- Reference relevant case law when appropriate (e.g., MWB Everton Park v Devcon Building Co, Southern Han, Dart Holdings, Galaxy Developments)
- Be clear about what will pass vs fail
- Use Australian legal terminology and spelling

When analyzing uploaded payment claims:
- Check each of the five requirements systematically
- Look for specific text that satisfies or fails each requirement
- Quote relevant excerpts from the document
- Give a clear pass/fail/warning for each requirement
- Provide specific recommendations for improvement

Be helpful but professional. If asked about topics outside the five requirements, politely redirect to what you can help with.
"""


def extract_text_from_pdf_bifact(pdf_file):
    """Extract text from PDF file for BIF Act analysis"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return f"Error reading PDF: {str(e)}"


def encode_image_bifact(image_bytes):
    """Encode image to base64 for GPT-4 Vision"""
    import base64
    try:
        image = Image.open(io.BytesIO(image_bytes))
        # Convert to RGB if necessary
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Resize if too large (max 2048x2048 for GPT-4 Vision)
        max_size = 2048
        if image.width > max_size or image.height > max_size:
            image.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)
        
        # Convert to base64
        buffered = io.BytesIO()
        image.save(buffered, format="JPEG")
        return base64.b64encode(buffered.getvalue()).decode('utf-8')
    except Exception as e:
        print(f"Error encoding image: {e}")
        return None


@app.post("/api/chat")
async def bifact_chat(
    message: str = Form(""),
    history: str = Form("[]"),
    file: Optional[UploadFile] = File(None)
):
    """
    BIF Act AI Assistant endpoint
    Handles both Q&A and document compliance checking
    """
    try:
        # Initialize OpenAI client
        openai_api_key = os.getenv("OPENAI_API_KEY")
        if not openai_api_key:
            raise HTTPException(status_code=500, detail="OpenAI API key not configured")
        
        client = OpenAI(api_key=openai_api_key)
        
        # Parse conversation history
        try:
            conversation_history = json.loads(history)
        except:
            conversation_history = []
        
        # Check if this is a document analysis request
        if file:
            # Read file content
            file_content = await file.read()
            filename = file.filename.lower()
            
            # Determine file type and extract content
            if filename.endswith('.pdf'):
                # Extract text from PDF
                document_text = extract_text_from_pdf_bifact(io.BytesIO(file_content))
                
                # Use GPT-4 for text analysis
                messages = [
                    {"role": "system", "content": BIF_ACT_SYSTEM_PROMPT + "\n\nAnalyze the following payment claim document and check compliance with the three requirements. For each requirement, determine if it PASSES, FAILS, or needs a WARNING. Provide specific quotes from the document and detailed explanations."},
                    {"role": "user", "content": f"Please analyze this payment claim for compliance:\n\n{document_text[:8000]}"}
                ]
                
                model = "gpt-4-turbo-preview"
                
            elif filename.endswith(('.png', '.jpg', '.jpeg')):
                # Handle image with GPT-4 Vision
                base64_image = encode_image_bifact(file_content)
                
                if not base64_image:
                    raise HTTPException(status_code=400, detail="Failed to process image")
                
                messages = [
                    {"role": "system", "content": BIF_ACT_SYSTEM_PROMPT + "\n\nAnalyze the payment claim document image and check compliance with the three requirements. For each requirement, determine if it PASSES, FAILS, or needs a WARNING. Provide specific quotes from the document and detailed explanations."},
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "Please analyze this payment claim document for compliance with the three BIF Act requirements (identify works, state amount, request payment)."
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ]
                
                model = "gpt-4-vision-preview"
            else:
                raise HTTPException(status_code=400, detail="Unsupported file type. Please upload PDF, PNG, or JPG.")
            
            # Add structured output request
            messages.append({
                "role": "system",
                "content": """Respond with a JSON object in this exact format:
{
  "requirements": [
    {
      "title": "Identify the Construction Works",
      "status": "pass|fail|warning",
      "explanation": "Detailed explanation of why it passes/fails/needs attention",
      "quote": "Relevant quote from the document (if available)"
    },
    {
      "title": "State the Claimed Amount",
      "status": "pass|fail|warning",
      "explanation": "Detailed explanation",
      "quote": "Relevant quote"
    },
    {
      "title": "Request Payment",
      "status": "pass|fail|warning",
      "explanation": "Detailed explanation",
      "quote": "Relevant quote"
    },
    {
      "title": "Claim for a Valid Reference Date",
      "status": "pass|fail|warning",
      "explanation": "Detailed explanation",
      "quote": "Relevant quote"
    },
    {
      "title": "Not Be for Unlicensed Building Work",
      "status": "pass|fail|warning",
      "explanation": "Detailed explanation - check if claimant appears to be properly licensed",
      "quote": "Relevant quote (e.g., licence number if shown)"
    }
  ]
}"""
            })
            
            # Call OpenAI API
            response = client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=0.3,
                max_tokens=2000
            )
            
            assistant_message = response.choices[0].message.content
            
            # Try to parse JSON response
            try:
                json_start = assistant_message.find('{')
                json_end = assistant_message.rfind('}') + 1
                json_str = assistant_message[json_start:json_end]
                compliance_result = json.loads(json_str)
            except:
                # If JSON parsing fails, return raw response
                compliance_result = {
                    "requirements": [
                        {
                            "title": "Analysis",
                            "status": "warning",
                            "explanation": assistant_message,
                            "quote": None
                        }
                    ]
                }
            
            return JSONResponse({
                "message": assistant_message,
                "compliance_check": compliance_result
            })
            
        else:
            # This is a Q&A request
            messages = [
                {"role": "system", "content": BIF_ACT_SYSTEM_PROMPT}
            ]
            
            # Add conversation history (limit to last 10 messages)
            for msg in conversation_history[-10:]:
                messages.append(msg)
            
            # Add current message
            messages.append({"role": "user", "content": message})
            
            # Call OpenAI API
            response = client.chat.completions.create(
                model="gpt-4-turbo-preview",
                messages=messages,
                temperature=0.7,
                max_tokens=2000
            )
            
            assistant_message = response.choices[0].message.content
            
            return JSONResponse({
                "message": assistant_message,
                "compliance_check": None
            })
            
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error in BIF Act chat endpoint: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to process request: {str(e)}")


@app.get("/api/chat/health")
async def bifact_chat_health():
    """Health check for BIF Act chatbot"""
    openai_api_key = os.getenv("OPENAI_API_KEY")
    return JSONResponse({
        "status": "healthy",
        "openai_configured": bool(openai_api_key)
    })




# ---------- serve frontend with clean URLs ----------
@app.get("/{path_name:path}")
async def serve_html_page(path_name: str):
    # --- START OF NEW FIX ---
    # Ignore paths that are clearly intended for API endpoints
    # In the serve_html_page function
    api_prefixes = [
        "api/", "admin/", "check-adjudicator-access/", "create-payment-intent/", 
        "purchase-register", "purchase-login", "purchase-me", "update-profile", 
        "my-purchases", "api/chat"
    ]
    if any(path_name.startswith(prefix) for prefix in api_prefixes):
        raise HTTPException(status_code=404, detail="API endpoint not found")


    if not path_name:
        path_name = "index"


    file_path = os.path.join(SITE_DIR, f"{path_name}.html")
    
    if os.path.abspath(file_path).startswith(os.path.abspath(SITE_DIR)) and os.path.exists(file_path) and not os.path.isdir(file_path):
        return FileResponse(file_path)
    
    static_file_path = os.path.join(SITE_DIR, path_name)
    if os.path.abspath(static_file_path).startswith(os.path.abspath(SITE_DIR)) and os.path.exists(static_file_path) and not os.path.isdir(static_file_path):
        return FileResponse(static_file_path)

    # Fallback for SPA-like behavior / non-html pages
    index_path = os.path.join(SITE_DIR, "index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path)

    raise HTTPException(status_code=404, detail="Not Found")
