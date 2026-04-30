"""
Extract every section in the annotated BIF Act DOCX into its own file.

Output: bif_guide_build/v3/source/annotated/section_NNN.txt

Each output file contains the full annotated commentary for one section,
including statutory text (verbatim, copied from DOCX) and the author's
commentary, with style tags preserved so the downstream agents can tell
legislation paragraphs from commentary paragraphs and from judicial extracts.

The chapter heading is included as a header line at the top of each file
so the agent knows which chapter the section sits in.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parent.parent.parent
DOCX_PATH = REPO / "legal_corpus" / "annotated_bif_act" / "v29 Annotated BIFSOPA (March 2026) (Clean)-1.docx"
OUT_DIR = Path(__file__).resolve().parent / "source" / "annotated"

SECTION_HEADER_STYLE = "2 Com-BIFSOPA Heading 1"
CHAPTER_HEADER_STYLE = "2a BIFSOPA Chapter Heading"
# section header pattern: "SECTION 67 – Meaning of reference date" or "SECTION 34A – ..."
SECTION_RE = re.compile(r"^\s*SECTION\s+(\d{1,3}[A-Z]?)\s*[–\-]\s*(.+?)\s*$", re.IGNORECASE)


def render_paragraph(p) -> str:
    style = p.style.name
    text = p.text.replace("\t", "    ")
    if not text.strip():
        return ""
    return f"[{style}] {text}"


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(DOCX_PATH))

    # First pass: locate every section header and the surrounding chapter.
    boundaries: list[tuple[str, str, str, int]] = []  # (chapter_header, section_id, section_title, para_index)
    current_chapter = ""
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == CHAPTER_HEADER_STYLE:
            current_chapter = p.text.strip()
            continue
        if p.style.name == SECTION_HEADER_STYLE:
            m = SECTION_RE.match(p.text.strip())
            if m:
                boundaries.append((current_chapter, m.group(1), m.group(2), i))

    # Second pass: each section's content runs to (next section start - 1) or end of doc.
    written = 0
    for idx, (chapter, sid, title, start) in enumerate(boundaries):
        end = boundaries[idx + 1][3] if idx + 1 < len(boundaries) else len(doc.paragraphs)

        # Build content with style tags preserved.
        lines = [
            f"# Annotated BIF Act source — Section {sid}",
            f"# Chapter: {chapter}",
            f"# Section title: {title}",
            f"# DOCX paragraphs: {start}-{end - 1}",
            "",
        ]
        for p in doc.paragraphs[start:end]:
            rendered = render_paragraph(p)
            if rendered:
                lines.append(rendered)
            else:
                lines.append("")

        # Filename uses zero-padded numeric base + optional letter suffix
        m = re.match(r"^(\d+)([A-Z]?)$", sid)
        if m:
            num, letter = m.group(1), m.group(2)
            slug = f"{int(num):03d}{letter}"
        else:
            slug = sid
        out = OUT_DIR / f"section_{slug}.txt"
        out.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
        written += 1

    print(f"wrote {written} per-section files to {OUT_DIR.relative_to(REPO)}/")


if __name__ == "__main__":
    main()
