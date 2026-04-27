"""
Extract Chapter 3 of the annotated BIF Act DOCX into 15 named source chunks.

Each output chunk file is a tagged plain-text dump of the DOCX paragraphs
falling within the requested section range. Every paragraph is prefixed with
its style name in square brackets so the downstream subagent can tell
legislation paragraphs (CDI styles + "1 BIFSOPA Heading") apart from
commentary paragraphs (Com-BIFSOPA styles + judicial extracts).

Run from the repo root:
    python3 bif_guide_build/extract_chunks.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parent.parent
DOCX_PATH = REPO / "legal_corpus" / "annotated_bif_act" / "v29 Annotated BIFSOPA (March 2026) (Clean)-1.docx"
OUT_DIR = Path(__file__).resolve().parent / "source_chunks"

# Chunk map: filename -> (sections, optional scoping note)
CHUNKS: list[tuple[str, list[int], str]] = [
    ("chunk_01_application.txt", [61, 62, 63], ""),
    ("chunk_02_definitions_construction_contract.txt", [64],
     "FOCUS: only the 'construction contract' definition within s64. "
     "Other defined terms in s64 are out of scope for this page."),
    ("chunk_03_construction_work_rgs.txt", [65, 66], ""),
    ("chunk_04_reference_date.txt", [67, 70], ""),
    ("chunk_05_payment_claim.txt", [68, 75],
     "FOCUS: s68 in full, plus s75(1)-(5) only. Later sub-sections of s75 "
     "(final claims after termination etc.) are out of scope for this page."),
    ("chunk_06_payment_schedule.txt", [69, 76], ""),
    ("chunk_07_consequences_no_schedule.txt", [77, 78], ""),
    ("chunk_08_adjudication_application.txt", [79], ""),
    ("chunk_09_adjudicator_appointment.txt", [80, 81], ""),
    ("chunk_10_adjudication_response.txt", [82, 83], ""),
    ("chunk_11_adjudication_procedures.txt", [84, 85, 86], ""),
    # NOTE: spec lists 's88, s87, s89' for chunk 12 — the page is built around
    # s88 (the decision itself) with s87 (valuation in later applications) as
    # an upstream rule that feeds the decision and s89 (slip-rule corrections)
    # as post-decision housekeeping. The chunk file is ordered by section
    # number (s87 -> s88 -> s89) so the source reads naturally; the subagent
    # is told to lead the page with s88.
    ("chunk_12_adjudicators_decision.txt", [87, 88, 89],
     "PAGE STRUCTURE: lead with s88 (Adjudicator's decision — primary rule). "
     "Treat s87 (valuation in later adjudications) as upstream context that "
     "feeds the decision. Treat s89 (slip-rule corrections) as post-decision "
     "housekeeping. In this chunk file the sections are in numerical order "
     "(87, 88, 89) so the source reads naturally — re-order on the page."),
    ("chunk_13_enforcement.txt", [90, 91, 92, 93], ""),
    ("chunk_14_other_rights_suspension.txt", [94, 95, 96, 97, 98], ""),
    ("chunk_15_civil_proceedings_service.txt", [99, 100, 101, 102], ""),
]

SECTION_HEADER_STYLE = "2 Com-BIFSOPA Heading 1"
SECTION_HEADER_RE = re.compile(r"^SECTION\s+(\d{1,3})\s*[–\-]", re.IGNORECASE)


def find_section_boundaries(doc: Document) -> dict[int, tuple[int, int]]:
    """Return {section_number: (start_para_index, end_para_index_exclusive)}.

    Boundaries cover Chapter 3 only (s61 onwards, stopping at the start of
    Chapter 4). The SECTION header paragraph itself is included in the range.
    """
    starts: list[tuple[int, int]] = []  # (section_number, para_index)
    chapter_4_start: int | None = None

    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "2a BIFSOPA Chapter Heading" and p.text.strip().startswith("CHAPTER 4"):
            chapter_4_start = i
            break
        if p.style.name == SECTION_HEADER_STYLE:
            m = SECTION_HEADER_RE.match(p.text.strip())
            if m:
                n = int(m.group(1))
                if 61 <= n <= 102:
                    starts.append((n, i))

    if chapter_4_start is None:
        raise RuntimeError("Could not find CHAPTER 4 boundary")

    boundaries: dict[int, tuple[int, int]] = {}
    for idx, (n, start) in enumerate(starts):
        end = starts[idx + 1][1] if idx + 1 < len(starts) else chapter_4_start
        boundaries[n] = (start, end)
    return boundaries


def render_paragraph(p) -> str:
    style = p.style.name
    text = p.text.replace("\t", "    ")
    if not text.strip():
        return ""
    return f"[{style}] {text}"


def render_section(doc: Document, n: int, start: int, end: int) -> str:
    lines: list[str] = []
    lines.append("=" * 78)
    lines.append(f"SECTION {n}  (DOCX paragraphs {start}-{end - 1})")
    lines.append("=" * 78)
    for p in doc.paragraphs[start:end]:
        rendered = render_paragraph(p)
        if rendered:
            lines.append(rendered)
        else:
            lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def main() -> None:
    if not DOCX_PATH.exists():
        raise SystemExit(f"DOCX not found: {DOCX_PATH}")
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document(str(DOCX_PATH))
    boundaries = find_section_boundaries(doc)
    missing = [n for n in range(61, 103) if n not in boundaries]
    if missing:
        raise SystemExit(f"Missing sections in DOCX: {missing}")

    for filename, sections, note in CHUNKS:
        body_parts: list[str] = []
        header = [
            f"# {filename}",
            f"# Sections covered: {', '.join(f's{n}' for n in sections)}",
        ]
        if note:
            header.append(f"# {note}")
        header.append("")
        body_parts.append("\n".join(header))

        for n in sections:
            start, end = boundaries[n]
            body_parts.append(render_section(doc, n, start, end))

        out_path = OUT_DIR / filename
        out_path.write_text("\n".join(body_parts), encoding="utf-8")
        print(f"wrote {out_path.relative_to(REPO)}  ({out_path.stat().st_size:>7d} bytes)")


if __name__ == "__main__":
    main()
