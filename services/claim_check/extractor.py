"""Document text extraction for /claim-check.

Supported:
  - PDF   -> pdfplumber (primary), PyPDF2 fallback
  - DOCX  -> python-docx
  - XLSX  -> openpyxl (structured: sheet -> rows -> cells)
  - Plain text (pasted)

The engine also needs structured data (for example: which sheet holds
the claim total) — ``extract_rich`` returns both flat text for prompt
construction AND structured extras the caller can feed to downstream
tooling.
"""

from __future__ import annotations

import io
import logging
from typing import Any

log = logging.getLogger("claim_check.extractor")

MAX_TEXT_CHARS = 200_000
MAX_XLSX_SHEET_ROWS_FOR_PROMPT = 400  # per-sheet cap for the prompt text blob
MAX_XLSX_SHEET_COLS_FOR_PROMPT = 40


def extract_text(filename: str, content: bytes) -> str:
    """Back-compat wrapper: flat text only."""
    text, _ = extract_rich(filename, content)
    return text


def extract_rich(filename: str, content: bytes) -> tuple[str, dict[str, Any]]:
    """Return (flat_text, extras). extras may include structured XLSX data.

    extras keys (all optional):
      - xlsx_structure: {"sheets": [{"name": str, "rows": [[cell, ...], ...]}, ...]}
    """
    if not content:
        raise ValueError("Empty document")

    name = (filename or "").lower()
    extras: dict[str, Any] = {}

    if name.endswith(".pdf"):
        text, pdf_meta = _extract_pdf_rich(content)
        extras.update(pdf_meta)
    elif name.endswith(".docx"):
        text = _extract_docx(content)
    elif name.endswith((".xlsx", ".xlsm")):
        text, extras = _extract_xlsx(content)
    elif name.endswith(".txt"):
        text = content.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type: {filename!r}. Use PDF, DOCX, XLSX, or paste the text.")

    text = (text or "").strip()
    if not text:
        raise ValueError("No text could be extracted from this document.")

    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS] + "\n\n[...document truncated at {:,} chars...]".format(MAX_TEXT_CHARS)

    return text, extras


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _extract_pdf(content: bytes) -> str:
    text, _extras = _extract_pdf_rich(content)
    return text


def _extract_pdf_rich(content: bytes) -> tuple[str, dict[str, Any]]:
    """Extract and also return per-file metadata (page count, scanned hint).

    Pages with rotation metadata different to the majority are normalised
    by pdfplumber when extracting text — the downstream LLM therefore sees
    reading-order text regardless of /Rotate. A scanned PDF (very little
    extractable text) is flagged in the returned metadata so the frontend
    can surface a caveat banner.
    """
    meta: dict[str, Any] = {"pages": 0, "scanned": False}
    try:
        import pdfplumber
    except Exception as e:
        log.warning("pdfplumber import failed, falling back to PyPDF2: %s", e)
        text = _extract_pdf_pypdf2(content)
        # Can't know pages cheaply without the lib; PyPDF2 can — cheap-count.
        try:
            import PyPDF2
            r = PyPDF2.PdfReader(io.BytesIO(content))
            meta["pages"] = len(r.pages)
        except Exception:
            pass
        return text, meta

    parts: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            meta["pages"] = len(pdf.pages)
            for page in pdf.pages:
                t = page.extract_text() or ""
                if t:
                    parts.append(t)
    except Exception as e:
        log.warning("pdfplumber extraction failed (%s), falling back to PyPDF2", e)
        text = _extract_pdf_pypdf2(content)
        return text, meta

    text = "\n\n".join(parts).strip()
    if not text:
        text = _extract_pdf_pypdf2(content)

    # Scanned-PDF heuristic: chars-per-page very low.
    if meta["pages"] and len(text) / max(1, meta["pages"]) < 60:
        meta["scanned"] = True

    return text, meta


def _extract_pdf_pypdf2(content: bytes) -> str:
    import PyPDF2
    reader = PyPDF2.PdfReader(io.BytesIO(content))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n\n".join(parts).strip()


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _extract_docx(content: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(content))

    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------

def _extract_xlsx(content: bytes) -> tuple[str, dict[str, Any]]:
    """Return (flat_text_for_prompt, extras).

    extras["xlsx_structure"] = {
        "sheets": [{"name": str, "rows": [[cell_value_or_null, ...], ...], "named_ranges": [...]}]
    }
    """
    try:
        import openpyxl
    except Exception as e:
        raise ValueError(f"openpyxl not available on the server: {e}")

    wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True, read_only=True)

    named_ranges: list[dict[str, str]] = []
    try:
        for dn in wb.defined_names.definedName if hasattr(wb.defined_names, "definedName") else wb.defined_names:
            try:
                named_ranges.append({"name": dn.name, "value": str(dn.value or dn.attr_text or "")})
            except Exception:
                continue
    except Exception:
        pass

    sheets_out: list[dict[str, Any]] = []
    text_parts: list[str] = []

    for ws in wb.worksheets:
        rows: list[list[Any]] = []
        row_count = 0
        col_cap = 0
        for row in ws.iter_rows(values_only=True):
            cells = list(row)
            rows.append(cells)
            row_count += 1
            if len(cells) > col_cap:
                col_cap = len(cells)

        # Build a textual representation for the prompt, capped.
        text_parts.append(f"# Sheet: {ws.title} ({row_count} rows × {col_cap} cols)")
        for i, cells in enumerate(rows[:MAX_XLSX_SHEET_ROWS_FOR_PROMPT]):
            trimmed = cells[:MAX_XLSX_SHEET_COLS_FOR_PROMPT]
            line_parts: list[str] = []
            for c in trimmed:
                if c is None:
                    line_parts.append("")
                else:
                    line_parts.append(str(c))
            if any(p for p in line_parts):
                text_parts.append(" | ".join(line_parts))
        if row_count > MAX_XLSX_SHEET_ROWS_FOR_PROMPT:
            text_parts.append(f"[...{row_count - MAX_XLSX_SHEET_ROWS_FOR_PROMPT} more rows omitted from prompt text...]")

        sheets_out.append({
            "name": ws.title,
            "rows": [[_cell_jsonable(c) for c in r] for r in rows],
            "row_count": row_count,
            "col_count": col_cap,
        })

    try:
        wb.close()
    except Exception:
        pass

    extras = {
        "xlsx_structure": {
            "sheets": sheets_out,
            "named_ranges": named_ranges,
        }
    }
    return "\n".join(text_parts).strip(), extras


def _cell_jsonable(v: Any) -> Any:
    import datetime as _dt
    if v is None:
        return None
    if isinstance(v, (str, int, float, bool)):
        return v
    if isinstance(v, (_dt.date, _dt.datetime)):
        return v.isoformat()
    return str(v)
