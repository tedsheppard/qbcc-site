"""DOCX export for Contract Assist drafts.

Produces a Word document with:
  - First-page DRAFT banner + disclaimer
  - "DRAFT — NOT FOR SERVICE" watermark behind every page (rotated VML)
  - Body content (the LLM's drafted text), with markdown-ish formatting
  - Disclaimer footer on every page
  - Citations footer at the end

python-docx doesn't expose watermark directly — we inject the underlying
VML XML into each section header. This is the documented hack for adding
true page watermarks via python-docx.

Streamed to the browser; never written to disk.
"""

from __future__ import annotations

import io
import logging
import re
from datetime import datetime, timezone
from typing import Any

log = logging.getLogger("contract_assist.draft_exporter")

WATERMARK_TEXT = "DRAFT — NOT FOR SERVICE"
DISCLAIMER_FOOTER = (
    "This is a draft prepared by an AI tool based on the information provided. "
    "Review, adapt, and verify before using. Consider obtaining legal advice before "
    "serving any formal notice."
)


def build_docx(*, kind: str, content: str, citations_contract: list[str] | None = None,
               citations_bif: list[str] | None = None) -> bytes:
    """Return DOCX bytes for the supplied draft body."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement

    doc = Document()

    # ---------- page setup ----------
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # ---------- watermark on every page (VML in headers) ----------
    for section in doc.sections:
        _add_watermark_to_header(section, WATERMARK_TEXT)
        _add_disclaimer_footer(section, DISCLAIMER_FOOTER)

    # ---------- styles ----------
    styles = doc.styles
    base = styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(11)

    # ---------- first-page header ----------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DRAFT — not for service")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(kind or "Draft document")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)

    # Disclaimer box (bordered paragraph).
    disc = doc.add_paragraph()
    disc.paragraph_format.space_before = Pt(8)
    disc.paragraph_format.space_after = Pt(14)
    r = disc.add_run(
        "This is a draft prepared by an AI tool based on the information provided. "
        "You must review, adapt, and verify this draft before using it. Consider "
        "obtaining legal advice before serving any formal notice."
    )
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x7A, 0x4A, 0x00)
    _shade_paragraph(disc, "FFFBEB")
    _border_paragraph(disc, "FDE68A")

    # ---------- body ----------
    body = _strip_draft_envelope(content or "")
    for para in body.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        _add_markdown_paragraph(doc, para)

    # ---------- citations block ----------
    if citations_contract or citations_bif:
        doc.add_paragraph()
        h = doc.add_paragraph()
        hr = h.add_run("Key references used in this draft")
        hr.bold = True
        hr.font.size = Pt(11)
        if citations_contract:
            for c in citations_contract:
                pp = doc.add_paragraph()
                pp.style = doc.styles["List Bullet"]
                pp.add_run(c)
        if citations_bif:
            for c in citations_bif:
                pp = doc.add_paragraph()
                pp.style = doc.styles["List Bullet"]
                pp.add_run(c)

    # ---------- export ----------
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_DRAFT_RULE_RE = re.compile(r"^\s*─{6,}\s*$|^\s*-{6,}\s*$")


def _strip_draft_envelope(text: str) -> str:
    """Remove the LLM's DRAFT envelope rules and the trailing disclaimer block —
    we render those ourselves with proper Word styling."""
    if not text:
        return ""
    # Drop the first DRAFT header line if present.
    lines = text.splitlines()
    while lines and (
        not lines[0].strip()
        or _DRAFT_RULE_RE.match(lines[0])
        or "DRAFT — not for service" in lines[0]
    ):
        lines.pop(0)
    # Drop everything after the closing rule (the trailing disclaimer + key refs).
    out_lines: list[str] = []
    for ln in lines:
        if _DRAFT_RULE_RE.match(ln):
            break
        out_lines.append(ln)
    return "\n".join(out_lines).strip()


def _add_markdown_paragraph(doc, text: str) -> None:
    """Very light markdown rendering — bold (**...**), italic (*...*), bullets,
    headings (#). Not a full markdown engine — just enough to make drafts look
    presentable in Word."""
    text = text.strip()
    if not text:
        return
    if text.startswith("# "):
        p = doc.add_paragraph()
        r = p.add_run(text[2:].strip())
        r.bold = True
        r.font.size = 14  # type: ignore
        return
    if text.startswith("## "):
        p = doc.add_paragraph()
        r = p.add_run(text[3:].strip())
        r.bold = True
        return
    if text.startswith("- ") or text.startswith("* "):
        for line in text.splitlines():
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, line.lstrip("-* ").strip())
        return
    if re.match(r"^\d+\.\s", text):
        for line in text.splitlines():
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, re.sub(r"^\d+\.\s+", "", line).strip())
        return

    p = doc.add_paragraph()
    _add_inline_runs(p, text.replace("\n", " "))


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITAL_RE = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")


def _add_inline_runs(paragraph, text: str) -> None:
    """Split text on **bold** and *italic* markers, adding runs accordingly."""
    # Build a simple token list.
    pos = 0
    tokens: list[tuple[str, str]] = []
    while pos < len(text):
        b = _BOLD_RE.search(text, pos)
        i = _ITAL_RE.search(text, pos)
        nxt = None
        if b and i:
            nxt = b if b.start() <= i.start() else i
        else:
            nxt = b or i
        if not nxt:
            tokens.append(("text", text[pos:]))
            break
        if nxt.start() > pos:
            tokens.append(("text", text[pos:nxt.start()]))
        if nxt is b:
            tokens.append(("bold", b.group(1)))
            pos = b.end()
        else:
            tokens.append(("ital", i.group(1)))
            pos = i.end()

    for kind, value in tokens:
        run = paragraph.add_run(value)
        if kind == "bold":
            run.bold = True
        elif kind == "ital":
            run.italic = True


def _add_watermark_to_header(section, text: str) -> None:
    """Inject a VML watermark shape into the section's primary header."""
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
    from lxml import etree

    header = section.header
    # Ensure the header has at least one paragraph.
    if not header.paragraphs:
        header.add_paragraph()
    p = header.paragraphs[0]
    # Build the VML watermark XML.
    xml = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        '     xmlns:v="urn:schemas-microsoft-com:vml" '
        '     xmlns:o="urn:schemas-microsoft-com:office:office" '
        '     xmlns:w10="urn:schemas-microsoft-com:office:word">'
        '  <w:pict>'
        f'    <v:shape id="SopalAssistWatermark" o:spid="_x0000_s1026" type="#_x0000_t136" '
        '            style="position:absolute;margin-left:0;margin-top:0;width:600pt;height:80pt;'
        '                   z-index:-251658240;mso-position-horizontal:center;'
        '                   mso-position-horizontal-relative:margin;'
        '                   mso-position-vertical:center;mso-position-vertical-relative:margin;'
        '                   rotation:-30" fillcolor="#dc2626" stroked="f">'
        '      <v:fill opacity="0.18"/>'
        f'      <v:textpath style="font-family:&quot;Calibri&quot;;font-size:1pt;font-weight:bold" '
        f'                  string="{text}"/>'
        '    </v:shape>'
        '  </w:pict>'
        '</w:r>'
    )
    fragment = etree.fromstring(xml)
    p._p.append(fragment)


def _add_disclaimer_footer(section, text: str) -> None:
    """Slim disclaimer line on every page footer."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    footer = section.footer
    if not footer.paragraphs:
        footer.add_paragraph()
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)


def _shade_paragraph(paragraph, hex_color: str) -> None:
    """Apply a background shade to a paragraph (used for the disclaimer box)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _border_paragraph(paragraph, hex_color: str) -> None:
    """Add a single-line border around a paragraph (for the disclaimer box)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "8")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), hex_color)
        pBdr.append(b)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Citation extraction — pulls [clause X] / [s N BIF Act] mentions out of the
# draft body so the DOCX appendix lists exactly what the model cited.
# ---------------------------------------------------------------------------

def extract_citations(text: str) -> tuple[list[str], list[str]]:
    """Return (contract_citations, bif_citations) deduplicated, in first-seen order."""
    contract: list[str] = []
    bif: list[str] = []
    seen_c: set[str] = set()
    seen_b: set[str] = set()
    for m in re.finditer(r"\[(?:clause|cl\.?)\s+(\d+(?:\.\d+)*[a-z]?)\]", text or "", re.I):
        c = f"clause {m.group(1)}"
        if c not in seen_c:
            seen_c.add(c)
            contract.append(c)
    for m in re.finditer(r"\[s\s+(\d+(?:\(\d+\))*(?:\([a-z]\))?(?:\([ivx]+\))?)\s+BIF\s+Act\]", text or "", re.I):
        b = f"s {m.group(1)} BIF Act"
        if b not in seen_b:
            seen_b.add(b)
            bif.append(b)
    return contract, bif
